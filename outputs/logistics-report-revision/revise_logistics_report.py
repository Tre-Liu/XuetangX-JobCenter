from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor
from docx.table import Table
from docx.text.paragraph import Paragraph


BASE = Path("/Users/liuhongzhe/Documents/专业建设/outputs/logistics-report-revision")
INPUT = BASE / "original.docx"
OUTPUT = BASE / "附件1-现代物流管理专业2026级职业本科人才培养方案修订调研报告26.06.11-标黄替换修订版.docx"
RED = RGBColor(192, 0, 0)


def child_text(child):
    texts = []
    for t in child.iter(qn("w:t")):
        if t.text:
            texts.append(t.text)
    return "".join(texts).strip()


def find_child(body, needle):
    for i, child in enumerate(body):
        if child_text(child) == needle:
            return i
    raise ValueError(f"Cannot find body child: {needle}")


def delete_between(body, start, end):
    for _ in range(end - start):
        body.remove(body[start])


def set_run_font(run, bold=False, color=None, size=10.5):
    if color is not None:
        run.font.color.rgb = color
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.bold = bold


def set_run_red(run, bold=False):
    set_run_font(run, bold=bold, color=RED)


def insert_para_after_plain(doc, anchor, text, bold=False, size=10.5, align=None):
    new_p = OxmlElement("w:p")
    anchor.addnext(new_p)
    para = Paragraph(new_p, doc._body)
    if align is not None:
        para.alignment = align
    run = para.add_run(text)
    set_run_font(run, bold=bold, size=size)
    return new_p


def insert_para_after(doc, anchor, text, style_name=None, bold=False, align=None):
    new_p = OxmlElement("w:p")
    anchor.addnext(new_p)
    para = Paragraph(new_p, doc._body)
    if style_name:
        try:
            para.style = style_name
        except Exception:
            pass
    if align is not None:
        para.alignment = align
    run = para.add_run(text)
    set_run_red(run, bold=bold)
    return new_p


def insert_table_after(doc, anchor, rows, style=None):
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    if style:
        try:
            table.style = style
        except Exception:
            pass
    for r_idx, row in enumerate(rows):
        for c_idx, value in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(value)
            set_run_red(run, bold=(r_idx == 0))
            p.paragraph_format.space_after = Pt(0)
            if r_idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tbl = table._tbl
    doc._body._body.remove(tbl)
    anchor.addnext(tbl)
    return tbl


def clone_table_style(doc, heading_text):
    body = doc._body._body
    idx = find_child(body, heading_text)
    for child in body[idx + 1 :]:
        if child.tag.endswith("tbl"):
            return Table(child, doc)._tbl.tblPr
        if child.tag.endswith("p") and child_text(child).startswith("五、"):
            break
    return None


section1_paragraphs = [
    ("（一）调研目标", True),
    (
        "本次修订以国家统计局、交通运输部、国家邮政局公开统计和辽宁、锦州区域产业规划为主要依据，重点核验现代物流管理专业在智慧仓储、运输调度、供应链计划、跨境与冷链物流、电商履约、物流数据分析等岗位上的真实需求。最新可核验的全国口径显示，2024年交通运输、仓储和邮政业增加值为59232亿元，同比增长7.0%；全年货物运输总量578.3亿吨，港口货物吞吐量176亿吨，港口集装箱吞吐量33200万标准箱，快递业务量1751亿件，电子商务交易额464091亿元、实物商品网上零售额127878亿元。这些数据说明，专业修订应从传统仓储运输管理，转向“数智运营+供应链协同+跨境与绿色物流”的复合能力培养。",
        False,
    ),
    ("二、总体调研目标", True),
    (
        "对接职业本科层次定位与现代物流产业升级要求，形成“岗位群识别—典型工作任务拆解—核心能力归纳—课程与实训项目调整—证书与竞赛融通”的闭环。调研重点不是泛泛描述物流行业，而是明确学生毕业3年内可进入的初始岗位、3—5年可发展的主管岗位，以及继续成长所需的数据、系统、流程和项目管理能力。",
        False,
    ),
    ("核心细分目标", True),
    (
        "（一）政策与定位目标：落实职业本科人才培养要求，服务辽宁沿海经济带、锦州港口型国家物流枢纽、东北陆海新通道和县域商业体系建设，突出智慧物流、供应链运营、跨境电商物流、冷链物流和绿色低碳物流方向。",
        False,
    ),
    (
        "（二）行业与岗位目标：以企业真实岗位为牵引，重点关注仓储运营专员、运输调度员、供应链计划专员、物流数据分析专员、跨境物流/货代操作、冷链物流运营、电商履约运营、物流项目执行等岗位，提取WMS/TMS/OMS系统操作、库存控制、线路优化、异常处理、数据看板、成本核算、客户沟通、质量安全与低碳合规等能力点。",
        False,
    ),
    (
        "（三）课程与教学目标：将《智慧仓储与配送管理》《物流运输管理》《供应链管理》《物流数字化管理》《国际货运代理》《冷链物流管理》《物流成本与绩效管理》等课程与岗位任务重新对齐，增加数据分析、业务流程建模、智能设备协同、绿色包装与碳核算、跨境单证与关务协同等训练内容。",
        False,
    ),
    (
        "（四）实践与质量目标：围绕京东、顺丰、港口物流、跨境电商和区域制造企业供应链场景，建设可评价的实训项目包；以毕业生就业对口率、岗位任务胜任度、企业满意度、证书获取与竞赛成果、课程项目产出作为培养方案迭代依据。",
        False,
    ),
    ("（二）调研对象", True),
    (
        "调研对象建议扩大为四类：一是辽宁省内同类院校与职业本科试点专业，重点比较课程体系、实训条件、岗课赛证融通做法；二是京东物流、顺丰、港口/货代、制造企业供应链、冷链与电商履约企业，重点采集岗位职责、招聘要求和系统工具；三是在校生、实习生和毕业生，重点验证课程适配度、岗位转化路径和能力短板；四是政府与行业公开数据，包括国家统计局、交通运输部、国家邮政局、辽宁省及锦州市有关物流、邮政快递、港口、多式联运、电子商务和数字经济的公开统计。",
        False,
    ),
    ("（三）调研方式", True),
]

section1_table = [
    ["调研对象", "调研方法", "形成材料"],
    [
        "同类院校与专业",
        "培养方案比对、课程标准比对、公开资料收集、访谈交流",
        "课程结构对标表、实训条件差距清单、职业本科定位参考",
    ],
    [
        "京东、顺丰、港口/货代、冷链、电商履约、制造供应链企业",
        "企业走访、线上座谈、岗位说明书采集、招聘信息样本分析",
        "岗位群清单、典型任务清单、系统工具与证书需求清单",
    ],
    [
        "在校生、实习生、毕业生",
        "问卷、访谈、就业跟踪、课程满意度调查",
        "课程适配度分析、就业岗位分布、能力短板与实训改进建议",
    ],
    [
        "政策与行业数据",
        "国家统计公报、交通运输和邮政行业公报、辽宁/锦州规划与统计资料研读",
        "产业发展依据、区域服务定位、数据口径说明",
    ],
]

section2_paragraphs = [
    ("（一）专业行业面向", True),
    (
        "本专业面向交通运输、仓储和邮政业及现代供应链服务业，重点覆盖智慧仓储与配送、运输组织与调度、供应链计划与采购协同、电商履约与即时配送、国际物流与跨境电商物流、冷链物流、港口与多式联运、制造企业生产物流和物流数据运营等领域。全国数据表明，2024年货物运输总量578.3亿吨、货物运输周转量261948.1亿吨公里，港口货物吞吐量176亿吨，外贸货物吞吐量54亿吨，快递业务量1751亿件；行业规模与业务复杂度同步提升，要求毕业生既懂物流流程，又能使用信息系统和数据工具解决现场问题。",
        False,
    ),
    ("（二）专业职业面向", True),
    (
        "核心岗位群建议调整为六类：1.智慧仓储与配送运营岗位群，面向仓储运营专员、库内规划员、配送运营专员；2.运输调度与多式联运岗位群，面向运输调度员、运力管理专员、港口/铁路/公路联运操作；3.供应链计划与采购协同岗位群，面向供应链计划专员、采购执行专员、库存控制专员；4.电商与跨境物流岗位群，面向电商履约运营、跨境物流操作、国际货代与单证关务专员；5.冷链与绿色物流岗位群，面向冷链运营、温控质量管理、绿色包装与碳数据采集岗位；6.物流数据与项目运营岗位群，面向物流数据分析助理、WMS/TMS/OMS系统运营、物流项目执行与客户方案助理。",
        False,
    ),
    (
        "上述岗位的共同能力从“会执行流程”升级为“会用系统、会看数据、会处理异常、会协同客户与供应商、会控制成本与质量”。因此，人才培养方案中应提高数据分析、系统应用、流程优化、成本绩效、标准化与安全合规、跨境单证、绿色低碳等内容比重。",
        False,
    ),
    ("（三）服务区域", True),
    (
        "服务区域建议定位为“立足锦州、服务辽宁、辐射东北陆海新通道和东北亚经贸物流网络”。锦州港口型国家物流枢纽、辽宁沿海经济带、沈阳—大连综合物流枢纽、东北制造业基地和跨境电商业务共同构成专业服务面向。专业不宜只面向单一企业或单一仓储场景，而应覆盖港口物流、生产制造物流、商贸流通物流、农产品冷链、跨境电商物流和县域配送网络。",
        False,
    ),
    ("（四）行业发展背景", True),
    (
        "从全国趋势看，物流业正在从规模扩张转向效率、韧性和绿色低碳并重。2024年交通运输、仓储和邮政业固定资产投资同比增长5.9%，新增港口万吨级及以上码头泊位通过能力27106万吨/年；电子商务交易额464091亿元、实物商品网上零售额127878亿元，说明电商履约、仓配一体和供应链协同仍是需求增量。对职业本科毕业生而言，企业更看重其在现场系统、数据看板、库存策略、运输调度、异常处置、成本核算和项目落地方面的综合能力。",
        False,
    ),
    (
        "从辽宁与锦州场景看，港口型物流枢纽、多式联运、装备制造供应链、农产品与医药冷链、跨境电商和县域商业配送是最需要补强的人才场景。人才培养方案应把“仓储—运输—订单—库存—客户—成本—质量—数据”的全流程作为主线，形成区别于普通专科的系统化设计能力，也区别于普通本科的现场实施与技术技能优势。",
        False,
    ),
]

section3_paragraphs = [
    (
        "本专业岗位发展路径建议按“初始岗位—发展岗位—目标岗位”组织，突出毕业生入职即能承担系统操作和现场执行任务，3—5年能够承担班组管理、流程优化和客户项目任务，5年以上具备供应链方案、成本绩效和跨部门协同能力。",
        False,
    )
]

section3_table = [
    ["岗位层次", "典型核心岗位", "主要工作内容", "发展性任务"],
    [
        "初始岗位（毕业1—2年）",
        "仓储运营专员、运输调度员、电商履约运营、跨境物流操作、冷链运营助理、物流数据助理",
        "完成入出库、库存盘点、订单履约、车辆调度、单证处理、温控记录、WMS/TMS/OMS数据维护和异常跟进",
        "熟练掌握岗位SOP和系统操作，能独立处理一般异常，形成数据记录、服务意识和质量安全意识",
    ],
    [
        "发展岗位（毕业3—5年）",
        "仓储/运输主管、供应链计划专员、物流项目执行、货代/关务主管、冷链质量管理员",
        "组织班组作业，制定排班和运输计划，开展库存控制、供应商协同、项目进度跟踪、成本核算和客户沟通",
        "具备流程优化、KPI分析、跨部门协同、成本控制和项目复盘能力，能够带教新人并参与方案设计",
    ],
    [
        "目标岗位（毕业5年以上）",
        "物流运营经理、供应链主管、物流项目经理、区域仓配负责人、跨境/冷链业务负责人",
        "负责区域或项目整体运营，设计仓配与运输方案，统筹客户、供应商、系统、成本、质量与安全合规",
        "具备供应链系统思维、数据决策、商务沟通、风险管理和持续改善能力，能够推动企业物流数字化升级",
    ],
]

section4_table = [
    ["序号", "对应核心课程", "典型工作任务", "任务简要描述"],
    ["1", "智慧仓储与配送管理", "仓储作业流程设计与库位优化", "依据订单结构和SKU特征，完成收货、上架、拣选、复核、包装、出库流程设计，使用WMS进行库位、批次、盘点和异常管理。"],
    ["2", "物流运输管理", "运输方案设计与调度执行", "根据时效、成本、货物属性和服务承诺选择运输方式，完成车辆/运力调度、线路规划、在途跟踪、异常处理和运费核算。"],
    ["3", "供应链管理", "需求计划、库存控制与供应链协同", "基于销售预测、库存周转和供应商交付周期制定补货计划，设置安全库存，跟踪缺货、滞销、呆滞料和供应风险。"],
    ["4", "采购管理", "采购执行与供应商管理", "完成采购需求确认、询报价、下单、到货跟踪、验收入库、对账和供应商绩效记录，支持成本与质量改进。"],
    ["5", "物流数字化管理", "物流数据采集、清洗与运营看板制作", "从WMS/TMS/OMS、订单平台和设备端采集数据，完成指标口径整理、异常识别、可视化看板和运营改进建议。"],
    ["6", "物流信息系统规划与设计", "WMS/TMS/OMS系统配置与基础维护", "维护商品、库位、车辆、客户、运价和路由等主数据，处理系统单据、状态更新、权限配置和基础报表。"],
    ["7", "电商物流运营", "订单履约与大促保障", "组织电商订单拣货、打包、发运和退换货处理，制定大促期间人力、库位、运力、包装与客服异常预案。"],
    ["8", "国际货运代理/国际物流实务", "国际货运订舱、单证与关务协同", "完成订舱、报关报检资料准备、提单与运费核算、货物跟踪和跨境运输衔接，配合处理通关异常。"],
    ["9", "冷链物流管理", "冷链温控流程与质量追溯", "制定冷链仓储和运输方案，监控温湿度、时效、包装和交接记录，处理断链风险并形成追溯资料。"],
    ["10", "物流成本与绩效管理", "物流成本核算与运营KPI分析", "归集仓储、运输、包装、人工、异常赔付等成本，计算订单履约成本、库存周转、准时率、破损率等指标。"],
    ["11", "物流项目运营", "客户物流方案与项目落地", "根据客户业务流程设计仓配、运输或供应链服务方案，编制项目计划，跟踪上线、试运行、质量验收和复盘改进。"],
    ["12", "生产物流管理", "线边配送与生产物料保障", "依据生产计划组织物料配送、JIT/JIS补给、线边库存控制和异常响应，保障生产节拍与供应连续性。"],
    ["13", "物流金融与区块链技术", "供应链金融场景资料整理与风险识别", "围绕订单、仓单、运输单据和应收账款等资料，识别真实性、权属、履约和信用风险，支持供应链金融业务合规。"],
    ["14", "绿色物流与安全管理", "绿色包装、碳数据采集与安全合规", "执行绿色包装、循环载具、节能运输和安全生产要求，采集能耗、碳排、温控和安全检查数据，形成合规记录。"],
]

source_note = [
    ("数据来源说明（本次标黄内容替换依据）", True),
    (
        "1. 国家统计局《中华人民共和国2024年国民经济和社会发展统计公报》（2025-02-28）：交通运输、仓储和邮政业增加值、货运量、港口吞吐量、快递业务量、电子商务交易额、网上零售额、固定资产投资等数据。",
        False,
    ),
    (
        "2. 交通运输部《2024年交通运输行业发展统计公报》、国家邮政局《2024年邮政行业发展统计公报》及辽宁省、锦州市公开规划/统计资料：用于校核综合交通、邮政快递、港口枢纽、多式联运、区域服务定位等表述。",
        False,
    ),
    (
        "3. 企业与岗位口径参考京东、顺丰、港口/货代、冷链、电商履约和制造供应链等典型场景，将岗位能力归纳为系统操作、数据分析、现场组织、客户协同、成本绩效、质量安全和绿色低碳七类。",
        False,
    ),
]


def replace_section(doc, heading, next_heading, paragraphs=None, table_rows=None):
    body = doc._body._body
    h_idx = find_child(body, heading)
    n_idx = find_child(body, next_heading)
    delete_between(body, h_idx + 1, n_idx)
    anchor = body[h_idx]
    if paragraphs:
        for text, bold in paragraphs:
            anchor = insert_para_after(doc, anchor, text, style_name=None, bold=bold)
    if table_rows:
        anchor = insert_table_after(doc, anchor, table_rows)


def replace_heading_text_red(doc, old, new):
    body = doc._body._body
    idx = find_child(body, old)
    para = Paragraph(body[idx], doc._body)
    for run in para.runs:
        run.text = ""
    run = para.add_run(new)
    set_run_red(run, bold=True)


def replace_toc(doc):
    body = doc._body._body
    toc_idx = None
    for i, child in enumerate(body):
        if child.tag.endswith("sdt") and "目录" in child_text(child):
            toc_idx = i
            break
    if toc_idx is None:
        return
    anchor = body[toc_idx - 1]
    body.remove(body[toc_idx])
    toc_items = [
        "目录",
        "一、调研设计与实施",
        "（一）调研目标",
        "（二）调研对象",
        "（三）调研方式",
        "二、现代物流管理专业人才需求分析和预测",
        "（一）专业行业面向",
        "（二）专业职业面向",
        "（三）服务区域",
        "（四）行业发展背景",
        "三、确定岗位及发展性任务",
        "四、典型工作任务及简要描述",
        "五、企业实际需要的职业证书",
        "数据来源说明（本次标黄内容替换依据）",
        "附表1  2026级职业本科人才培养方案修订领导小组成员名单",
        "附表2  现代物流管理专业培养方案调研情况记录表",
    ]
    for n, item in enumerate(toc_items):
        anchor = insert_para_after_plain(doc, anchor, item, bold=(n == 0), size=12 if n == 0 else 10.5)


def main():
    doc = Document(INPUT)
    replace_toc(doc)

    replace_section(doc, "一、调研设计与实施", "二、现代物流管理专业人才需求分析和预测", section1_paragraphs, section1_table)
    replace_section(doc, "二、现代物流管理专业人才需求分析和预测", "确定岗位及发展性任务", section2_paragraphs)
    replace_heading_text_red(doc, "确定岗位及发展性任务", "确定岗位及发展性任务")
    replace_section(doc, "确定岗位及发展性任务", "四、典型工作任务及简要描述", section3_paragraphs, section3_table)
    replace_section(doc, "四、典型工作任务及简要描述", "五、企业实际需要的职业证书", None, section4_table)

    body = doc._body._body
    appendix_idx = find_child(body, "附表1")
    anchor = body[appendix_idx - 1]
    for text, bold in source_note:
        anchor = insert_para_after(doc, anchor, text, bold=bold)

    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
