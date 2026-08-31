import hashlib
import json
import logging
import re
from collections import Counter, defaultdict
from pathlib import Path

import pdfplumber
from docx import Document
from pypdf import PdfReader

from pipeline import consume_task_table, detect_task_table_columns, parse_source_metadata


DOCUMENTS = Path("/Users/liuhongzhe/Desktop/新双高人培/documents")
OUT_DIR = Path(__file__).resolve().parent / "data"


def sha256(path):
    digest = hashlib.sha256()
    with path.open("rb") as handle:
        for chunk in iter(lambda: handle.read(1024 * 1024), b""):
            digest.update(chunk)
    return digest.hexdigest()


def clean(value):
    return re.sub(r"\s+", "", value or "")


def infer_major(texts, page_index, fallback):
    nearby = "\n".join(texts[max(0, page_index - 8) : page_index + 1])
    patterns = (
        r"专业名称[（(](?:专业代码|代码)[）)]\s*([^\n（(]{2,30})[（(](\d{6})[）)]",
        r"专业名称(?:和代码)?\s*[:：]?\s*([^\n（(]{2,30})[（(](\d{6})[）)]",
        r"《([^《》\n]{2,30})》专业人才培养方案",
        r"([^\n]{2,30})专业人才培养方案[（(]普通",
    )
    result = dict(fallback)
    matches = []
    for pattern in patterns:
        matches.extend(re.finditer(pattern, nearby))
    if matches:
        match = max(matches, key=lambda item: item.start())
        name = clean(match.group(1)).strip("《》")
        if name and len(name) <= 30:
            result["major"] = name
        if match.lastindex and match.lastindex >= 2:
            result["major_code"] = match.group(2)
    return result


def candidate_pages(texts):
    candidates = set()
    for index, text in enumerate(texts):
        normalized = clean(text)
        if "典型工作任务" in normalized and any(k in normalized for k in ("岗位", "职业能力", "能力要求")):
            candidates.update(range(index, min(len(texts), index + 4)))
    return sorted(candidates)


def table_has_header(table):
    return any(detect_task_table_columns(row or []) for row in (table or [])[:6])


def extract_pdf(path, file_hash):
    base = parse_source_metadata(path)
    reader = PdfReader(str(path))
    texts = [(page.extract_text() or "") for page in reader.pages]
    candidates = candidate_pages(texts)
    records = []
    state = None
    last_page = None
    header_counter = Counter()

    with pdfplumber.open(path) as pdf:
        for page_index in candidates:
            if last_page is None or page_index - last_page > 1:
                state = None
            page = pdf.pages[page_index]
            tables = page.extract_tables() or []
            for table_index, table in enumerate(tables, start=1):
                has_header = table_has_header(table)
                if not has_header and not state:
                    continue
                if has_header:
                    for row in table[:6]:
                        if detect_task_table_columns(row or []):
                            header_counter[" | ".join(clean(cell) for cell in row or [])] += 1
                            break
                extracted, state = consume_task_table(table, state)
                if not extracted:
                    continue
                metadata = infer_major(texts, page_index, base)
                for record in extracted:
                    record.update(metadata)
                    record.update(
                        {
                            "source_path": str(path),
                            "source_sha256": file_hash,
                            "source_page": page_index + 1,
                            "source_locator": f"PDF第{page_index + 1}页/表{table_index}",
                        }
                    )
                    records.append(record)
            last_page = page_index
    return records, header_counter, len(reader.pages), candidates


def extract_docx(path, file_hash):
    base = parse_source_metadata(path)
    document = Document(path)
    records = []
    state = None
    headers = Counter()
    for table_index, table in enumerate(document.tables, start=1):
        rows = [[cell.text for cell in row.cells] for row in table.rows]
        if table_has_header(rows):
            for row in rows[:6]:
                if detect_task_table_columns(row or []):
                    headers[" | ".join(clean(cell) for cell in row or [])] += 1
                    break
        elif not state:
            continue
        extracted, state = consume_task_table(rows, state)
        for record in extracted:
            record.update(base)
            record.update(
                {
                    "source_path": str(path),
                    "source_sha256": file_hash,
                    "source_page": None,
                    "source_locator": f"Word表{table_index}",
                }
            )
            records.append(record)
    return records, headers, len(document.paragraphs), list(range(len(document.tables)))


def main():
    logging.getLogger("pypdf").setLevel(logging.ERROR)
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    files = sorted([*DOCUMENTS.rglob("*.pdf"), *DOCUMENTS.rglob("*.docx")])
    by_hash = defaultdict(list)
    for path in files:
        by_hash[sha256(path)].append(path)

    all_records = []
    all_headers = Counter()
    audit = []
    for number, (file_hash, paths) in enumerate(by_hash.items(), start=1):
        path = paths[0]
        try:
            if path.suffix.lower() == ".pdf":
                records, headers, unit_count, candidates = extract_pdf(path, file_hash)
            else:
                records, headers, unit_count, candidates = extract_docx(path, file_hash)
            all_records.extend(records)
            all_headers.update(headers)
            audit.append(
                {
                    "sha256": file_hash,
                    "representative_path": str(path),
                    "duplicate_paths": [str(item) for item in paths],
                    "unit_count": unit_count,
                    "candidate_count": len(candidates),
                    "record_count": len(records),
                    "status": "ok",
                }
            )
        except Exception as exc:
            audit.append(
                {
                    "sha256": file_hash,
                    "representative_path": str(path),
                    "duplicate_paths": [str(item) for item in paths],
                    "record_count": 0,
                    "status": "error",
                    "error": f"{type(exc).__name__}: {exc}",
                }
            )
        if number % 20 == 0:
            print(f"processed {number}/{len(by_hash)} unique files; records={len(all_records)}", flush=True)

    deduped = []
    seen = set()
    for record in all_records:
        key = (
            record["source_sha256"],
            record["source_page"],
            record["source_job"],
            record["typical_task"],
            record["competency"],
        )
        if key not in seen:
            seen.add(key)
            deduped.append(record)

    (OUT_DIR / "source_records.json").write_text(json.dumps(deduped, ensure_ascii=False, indent=2), encoding="utf-8")
    (OUT_DIR / "source_audit.json").write_text(json.dumps(audit, ensure_ascii=False, indent=2), encoding="utf-8")
    (OUT_DIR / "header_variants.json").write_text(json.dumps(all_headers, ensure_ascii=False, indent=2), encoding="utf-8")
    summary = {
        "input_files": len(files),
        "unique_files": len(by_hash),
        "duplicate_files": len(files) - len(by_hash),
        "records": len(deduped),
        "source_jobs": len({record["source_job"] for record in deduped}),
        "typical_tasks": len({(record["source_job"], record["typical_task"]) for record in deduped}),
        "files_with_records": sum(1 for item in audit if item["record_count"]),
        "errors": sum(1 for item in audit if item["status"] == "error"),
    }
    (OUT_DIR / "summary.json").write_text(json.dumps(summary, ensure_ascii=False, indent=2), encoding="utf-8")
    print(json.dumps(summary, ensure_ascii=False, indent=2))


if __name__ == "__main__":
    main()
