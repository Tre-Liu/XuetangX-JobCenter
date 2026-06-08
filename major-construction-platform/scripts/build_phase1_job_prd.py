from __future__ import annotations

from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/liuhongzhe/Documents/专业建设/major-construction-platform")
OUT_DIR = ROOT / "outputs" / "prd"
OUT_PATH = OUT_DIR / "一期岗位中心产品功能需求文档.docx"
SHOT_DIR = OUT_DIR / "screenshots"

SCREENSHOTS = {
    "industry_layout": SHOT_DIR / "industry-layout-demo.jpeg",
    "job_analysis": SHOT_DIR / "job-analysis-demo.jpeg",
    "job_build_empty": SHOT_DIR / "job-build-empty-demo.jpeg",
    "job_build_list": SHOT_DIR / "job-build-list-demo.jpeg",
    "results_job_center": SHOT_DIR / "results-job-center-demo.jpeg",
}


def set_east_asia_font(run, east_asia: str = "Microsoft YaHei") -> None:
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.rFonts
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.append(r_fonts)
    r_fonts.set(qn("w:eastAsia"), east_asia)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color: str = "D9E2F3", size: str = "6") -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.first_child_found_in("w:tcBorders")
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    for edge in ("top", "left", "bottom", "right"):
        el = tc_borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            tc_borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), size)
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def set_paragraph_spacing(paragraph, before: int = 0, after: int = 0, line: float = 1.2) -> None:
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing = line


def add_run(paragraph, text: str, *, bold: bool = False, size: int = 11, color: str | None = None):
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    set_east_asia_font(run)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return run


def add_heading(doc: Document, text: str, level: int):
    p = doc.add_paragraph(style=f"Heading {level}")
    sizes = {1: 16, 2: 13, 3: 12}
    colors = {1: "1F4E79", 2: "2F75B5", 3: "1F4E79"}
    add_run(p, text, size=sizes[level], color=colors[level])
    set_paragraph_spacing(p, before={1: 16, 2: 12, 3: 8}[level], after={1: 8, 2: 6, 3: 4}[level])
    return p


def add_text(doc: Document, text: str, *, before: int = 0, after: int = 6):
    p = doc.add_paragraph()
    add_run(p, text, size=11)
    set_paragraph_spacing(p, before=before, after=after)
    return p


def add_bullets(doc: Document, items: list[str], *, before: int = 0, after: int = 4):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        add_run(p, item, size=11)
        set_paragraph_spacing(p, before=before, after=after)


def add_numbered(doc: Document, items: list[str], *, before: int = 0, after: int = 4):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        add_run(p, item, size=11)
        set_paragraph_spacing(p, before=before, after=after)


def add_caption(doc: Document, text: str):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = add_run(p, text, size=9, color="666666")
    run.italic = True
    set_paragraph_spacing(p, before=3, after=10)


def add_image(doc: Document, image_path: Path, caption: str, width: float = 6.1):
    if not image_path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(image_path), width=Inches(width))
    set_paragraph_spacing(p, before=6, after=0)
    add_caption(doc, caption)


def add_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str] | tuple[str, ...]],
    widths: list[float],
    *,
    header_fill: str = "EAF2F8",
    header_color: str = "1F1F1F",
):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    for idx, cell in enumerate(table.rows[0].cells):
        cell.width = Inches(widths[idx])
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        set_cell_border(cell, color="B8CCE4")
        set_cell_shading(cell, header_fill)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_run(p, headers[idx], bold=True, size=10, color=header_color)
        set_paragraph_spacing(p, before=3, after=3)

    for row in rows:
        tr = table.add_row().cells
        for idx, value in enumerate(row):
            tr[idx].width = Inches(widths[idx])
            tr[idx].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            set_cell_border(tr[idx])
            p = tr[idx].paragraphs[0]
            if idx == 0:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_run(p, value, size=10)
            set_paragraph_spacing(p, before=3, after=3)

    doc.add_paragraph()
    return table


def setup_document(doc: Document) -> None:
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(0.9)
    section.bottom_margin = Inches(0.8)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.paragraph_format.line_spacing = 1.2

    for style_name, size, color in (
        ("Heading 1", 16, "1F4E79"),
        ("Heading 2", 13, "2F75B5"),
        ("Heading 3", 12, "1F4E79"),
    ):
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=48, after=8)
    add_run(p, "岗位中心一期产品功能需求文档", bold=True, size=22, color="16365C")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(p, before=0, after=10)
    add_run(
        p,
        "覆盖范围：岗位数据调研-产业布局、岗位数据调研-岗位分析、岗位建设中心、查看成果页-岗位中心",
        size=11,
        color="666666",
    )

    meta_rows = [
        ("文档类型", "产品功能需求文档（研发交付版）"),
        ("版本", "V2.0"),
        ("适用阶段", "一期需求确认 / 研发实现 / 联调验收"),
        ("编制日期", str(date.today())),
        ("适用对象", "产品、前端、后端、测试、实施"),
        ("截图来源", "当前本地 demo 页面真实截图，不使用文档或表格截图"),
    ]
    add_table(doc, ["字段", "内容"], meta_rows, [1.5, 5.6], header_fill="F2F6FC")

    add_text(
        doc,
        "本文档用于把一期岗位中心相关需求从“汇报口径”细化到“可交付研发口径”。文档重点描述业务目标、页面结构、交互规则、空状态、异常状态和一期边界，不展开数据库设计、接口协议和技术架构方案。",
        before=6,
        after=10,
    )
    doc.add_page_break()


def add_module(
    doc: Document,
    title: str,
    purpose: str,
    entries: list[tuple[str, str]],
    req_rows: list[tuple[str, str, str]],
    state_rows: list[tuple[str, str, str]],
    flows: list[str],
    exclusions: list[str],
    screenshots: list[tuple[Path, str, float]],
) -> None:
    add_heading(doc, title, 1)
    add_text(doc, purpose)

    add_heading(doc, "1. 页面入口与对象", 2)
    add_table(doc, ["页面/对象", "说明"], entries, [1.9, 5.2], header_fill="EDF4FB")

    add_heading(doc, "2. 功能详细需求", 2)
    add_table(
        doc,
        ["需求编号", "功能点", "详细规则"],
        req_rows,
        [0.9, 1.7, 4.5],
        header_fill="E8F1FB",
    )

    add_heading(doc, "3. 状态与边界处理", 2)
    add_table(
        doc,
        ["场景", "触发条件", "处理规则"],
        state_rows,
        [1.5, 1.8, 3.8],
        header_fill="F7F9FC",
    )

    add_heading(doc, "4. 页面截图", 2)
    for image_path, caption, width in screenshots:
        add_image(doc, image_path, caption, width=width)

    add_heading(doc, "5. 关键业务流程", 2)
    add_numbered(doc, flows)

    add_heading(doc, "6. 一期不做", 2)
    add_bullets(doc, exclusions)
    doc.add_paragraph()


def build_doc() -> Path:
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    doc = Document()
    setup_document(doc)
    add_cover(doc)

    add_heading(doc, "一、文档目的", 1)
    add_text(
        doc,
        "一期目标是先打通“产业认知 -> 岗位判断 -> 岗位沉淀 -> 成果展示”的主链路，形成可演示、可理解、可继续扩展的岗位中心底座。研发实现时应优先保证页面链路、状态反馈和数据承接关系清晰可用。",
    )

    add_heading(doc, "二、一期范围", 1)
    add_table(
        doc,
        ["模块", "一期定位", "交付重点"],
        [
            ("岗位数据调研-产业布局", "看产业", "帮助用户理解产业链结构、区域分布、政策影响与代表企业。"),
            ("岗位数据调研-岗位分析", "看岗位", "帮助用户识别重点岗位、查看岗位画像，并判断需求趋势与新岗位方向。"),
            ("岗位建设中心", "做建设", "沉淀岗位图谱、岗位列表、岗位详情、课程关联和建设入口。"),
            ("查看成果页-岗位中心", "做展示", "将岗位建设结果转化为汇报和验收可用的成果页内容。"),
        ],
        [1.6, 1.3, 4.2],
    )

    add_heading(doc, "三、通用交互约定", 1)
    add_table(
        doc,
        ["主题", "规则"],
        [
            ("默认产业链", "一期默认展示“人工智能产业链”；页面顶部统一展示“当前产业链”选择入口，后续预留多产业链切换能力。"),
            ("页面组织", "岗位中心采用左侧菜单 + 右侧内容区结构；建设成果展示采用顶部导航切换栏目。"),
            ("空状态", "页面无数据时必须保留页面骨架、标题和后续动作按钮，不能直接留白。"),
            ("禁用态", "无法执行的确认按钮默认禁用，并通过计数或文案说明原因，如“已选 0 个”时禁用“添加到岗位建设中心”。"),
            ("异常态", "数据加载或操作失败时保留当前页面，显示失败提示和重试入口，不得直接清空用户已输入内容。"),
            ("缺失字段", "单条记录存在字段缺失时，页面保持可浏览，缺失项展示“暂无数据”或空标签，不因单字段缺失阻断整页展示。"),
        ],
        [1.5, 5.6],
        header_fill="F2F6FC",
    )

    add_module(
        doc=doc,
        title="四、模块需求一：岗位数据调研-产业布局",
        purpose="该模块用于回答“产业链怎么分、企业和岗位集中在哪些环节、区域资源在哪里、政策会对专业建设产生什么影响”四类问题，为后续岗位画像与岗位建设提供产业依据。",
        entries=[
            ("一级入口", "岗位中心 -> 岗位数据调研"),
            ("二级页签", "产业链图谱、区域产业分析、产业政策库、产业企业库"),
            ("页面对象", "AI总结条、图谱/地图/时间轴、企业卡片、政策详情弹层"),
        ],
        req_rows=[
            (
                "IL-01",
                "默认打开产业链图谱",
                "用户进入“岗位数据调研”后，默认展示“产业布局 / 产业链图谱”页面；页面头部展示当前产业链名称，左侧高亮当前子菜单。",
            ),
            (
                "IL-02",
                "产业链图谱展示",
                "页面需展示 AI 总结条、上游/中游/下游桑基式结构图、价值流判断、建设切入点、能力短板、代表企业和核心岗位建议，整体只读展示，不支持在线编辑。",
            ),
            (
                "IL-03",
                "区域产业分析",
                "区域页需展示全国企业区域热力图、覆盖省份/企业样本/重点城市/合作线索 KPI、TOP15 省份排行和区域合作方向；鼠标悬停省份时显示“省份名称 + 公司数量”。",
            ),
            (
                "IL-04",
                "产业政策库",
                "政策页需展示政策级别筛选、关键词搜索、政策时间轴、关键词云和年度趋势；点击任意政策后打开详情弹层，至少展示政策级别、发布时间、发布机构、政策摘要、对专业建设影响和建议转化任务。",
            ),
            (
                "IL-05",
                "产业企业库",
                "企业页需按企业卡片形式展示企业名称、技术方向、可对接岗位和合作建议；同时提供企业合作优先级表，供实施和教研团队快速识别合作对象。",
            ),
            (
                "IL-06",
                "页面定位",
                "本模块为研判页，不承担岗位维护、课程维护和成果发布职责；用户在本模块中的主要行为是浏览、筛选、对比和点击查看详情。",
            ),
        ],
        state_rows=[
            ("空数据", "当前产业链未配置产业图谱/政策/企业数据", "保留左侧菜单和页面标题，主内容区展示“暂无产业调研数据”，同时保留“当前产业链”入口和返回上级能力。"),
            ("筛选无结果", "政策关键词或级别筛选后无匹配数据", "政策列表区域展示“暂无匹配政策”，并提供“清空筛选”动作；不隐藏页面其余模块。"),
            ("地图悬停", "鼠标未停留在省份区域", "区域地图 tooltip 不显示，页面不应出现空白浮层或报错。"),
            ("详情异常", "政策详情/企业详情单条数据缺失", "弹层仍可打开，缺失字段展示“暂无数据”；不因单条详情异常导致整页不可用。"),
            ("加载失败", "产业数据接口失败或超时", "展示失败提示与“重试”按钮；已展示的左侧菜单、页签和页面标题不消失。"),
        ],
        flows=[
            "用户从“岗位中心”进入“岗位数据调研”，系统默认展示“产业布局 / 产业链图谱”。",
            "用户切换到区域产业分析、政策库或企业库，查看不同维度的产业依据。",
            "用户在政策页点击某条政策，打开详情弹层确认政策影响与建议转化任务。",
            "用户基于产业结论，继续进入“岗位分析”模块识别需要建设的重点岗位。",
        ],
        exclusions=[
            "一期不做多产业链并行对比。",
            "一期不做复杂报表导出和批量下载。",
            "一期不做政策库、企业库后台维护台账，只要求前台可浏览、可查看详情。",
        ],
        screenshots=[
            (SCREENSHOTS["industry_layout"], "图 1  真实 demo 截图：岗位数据调研-产业布局-产业链图谱", 6.25),
        ],
    )

    add_module(
        doc=doc,
        title="五、模块需求二：岗位数据调研-岗位分析",
        purpose="该模块用于回答“哪些岗位最重要、岗位需要什么能力、市场需求如何变化、未来还会出现哪些新岗位”四类问题，是岗位建设中心的直接输入来源。",
        entries=[
            ("一级入口", "岗位中心 -> 岗位数据调研"),
            ("二级页签", "岗位画像分析、招聘需求趋势、新岗位新技术预判"),
            ("页面对象", "搜索栏、热门岗位标签、岗位卡片、分页、详情弹层、图谱入口"),
        ],
        req_rows=[
            (
                "JA-01",
                "岗位画像分析首页",
                "页面需展示说明条、岗位搜索引擎、热门岗位搜索、岗位画像洞察和岗位卡片列表。默认展示第一页岗位卡片，并显示总条数和分页控件。",
            ),
            (
                "JA-02",
                "搜索与快捷触发",
                "用户可通过岗位名称、技能关键词或产业链环节发起搜索；点击热门岗位标签应直接定位到对应岗位详情或刷新结果列表。搜索框、热门标签和分页结果需保持统一口径。",
            ),
            (
                "JA-03",
                "岗位卡片信息",
                "卡片至少展示岗位名称、薪资区间、需求量、岗位层级、所属产业节点和 2-3 个能力关键词；点击卡片后打开岗位画像详情。",
            ),
            (
                "JA-04",
                "岗位画像详情",
                "详情弹层至少包含数据来源、薪资、学历要求、经验要求、岗位层级、职业路径、画像摘要、标签、三维能力分析、典型工作任务、推荐职业资格证书、对接专业、相关企业。",
            ),
            (
                "JA-05",
                "详情内二级查看",
                "详情弹层中的“岗位能力图谱”按钮可进入图谱页；点击证书或企业后打开二级详情弹层；关闭二级弹层后回到岗位画像详情，不丢失原浏览位置。",
            ),
            (
                "JA-06",
                "招聘需求趋势",
                "趋势页需展示近 12 月岗位需求、高频招聘岗位、平均薪资、企业样本、月度趋势图、技能热度条形图和热门岗位招聘明细表，定位为只读分析页。",
            ),
            (
                "JA-07",
                "新岗位新技术预判",
                "预判页需展示 AI 洞察、高潜技术方向、三年岗位机会矩阵、技术信号监测、新岗位库和人才培养方向建议表，用于支撑后续岗位建设优先级判断。",
            ),
        ],
        state_rows=[
            ("分页首尾", "位于第一页或最后一页", "上一页/下一页按钮需禁用，避免翻页越界；页码区仍显示总页数。"),
            ("搜索无结果", "搜索条件无匹配岗位", "结果区显示“暂无匹配岗位”，热门岗位标签继续保留，用户可一键切换其他关键词。"),
            ("缺失字段", "岗位详情部分字段为空", "详情页保留整体结构，空字段展示“暂无数据”；任务、证书、企业为空时各自展示独立空提示。"),
            ("弹层关闭", "用户点击关闭按钮或取消", "仅关闭当前层弹窗，返回上一层页面；不得重置用户此前的搜索结果和分页位置。"),
            ("加载失败", "岗位画像或趋势数据加载失败", "当前页展示错误提示和重试动作；已存在的筛选条件和搜索词保留。"),
        ],
        flows=[
            "用户进入“岗位分析 / 岗位画像分析”，先通过搜索或热门岗位定位重点岗位。",
            "用户点击岗位卡片，查看岗位画像详情、能力结构、典型工作任务和相关企业。",
            "用户进一步查看“招聘需求趋势”和“新岗位新技术预判”，判断岗位的市场热度与未来方向。",
            "用户将确认后的重点岗位带入“岗位建设中心”继续建设。",
        ],
        exclusions=[
            "一期不在本模块中直接编辑岗位标准。",
            "一期不做复杂多条件高级检索和批量导出。",
            "一期不展开算法模型、爬虫规则或数据清洗过程说明。",
        ],
        screenshots=[
            (SCREENSHOTS["job_analysis"], "图 2  真实 demo 截图：岗位数据调研-岗位分析-岗位画像分析", 6.25),
        ],
    )

    add_module(
        doc=doc,
        title="六、模块需求三：岗位建设中心",
        purpose="岗位建设中心是一期开工位最重的模块，用于把调研结果沉淀为岗位图谱、岗位列表、岗位详情和课程关系，形成后续成果展示和课程联动的基础数据。",
        entries=[
            ("一级入口", "岗位中心 -> 岗位建设中心"),
            ("主页面对象", "AI总结条、添加岗位、AI建岗、图谱区、岗位列表、空状态面板"),
            ("详情页对象", "基本信息、典型工作任务、岗位能力项、岗位能力图谱、适岗度评价要求"),
        ],
        req_rows=[
            (
                "JB-01",
                "首页双态展示",
                "岗位建设中心需区分“已有岗位数据”和“无岗位数据”两种状态。无数据时展示初始化说明、操作按钮和三步引导；有数据时展示 AI 总结、产业岗位课程图谱和岗位列表。",
            ),
            (
                "JB-02",
                "AI总结区",
                "有数据时汇总岗位数量、典型工作任务数量和岗位能力点数量；无数据时改为初始化引导文案，明确可通过“添加岗位 / 模版导入 / AI建岗”开始建设。",
            ),
            (
                "JB-03",
                "添加岗位弹窗",
                "点击“添加岗位”后打开弹窗，展示模版导入条、搜索框、候选岗位列表、已选数量和确认按钮。搜索支持按岗位名称、职业编码、产业节点和能力关键词过滤候选项。",
            ),
            (
                "JB-04",
                "候选岗位选择规则",
                "已在建设中心中的岗位候选卡片置灰禁用，并标记“已添加”；未添加岗位可多选，选中后候选卡片显示勾选状态，顶部“已选 N 个”实时更新。",
            ),
            (
                "JB-05",
                "确认添加规则",
                "当已选数量为 0 时，“添加到岗位建设中心”按钮禁用；用户确认添加后，岗位应进入岗位列表，同时刷新 AI 总结和产业岗位课程图谱。",
            ),
            (
                "JB-06",
                "岗位列表与删除",
                "岗位列表以卡片展示岗位名称、所属岗位群、典型任务数、岗位能力数和职业信息；删除岗位后，列表、统计和图谱同步刷新，删除仅取消当前专业关联，不删除岗位调研源数据。",
            ),
            (
                "JB-07",
                "岗位详情页",
                "点击岗位卡片进入岗位详情页。详情页左侧需提供“基本信息 / 典型工作任务 / 岗位能力项 / 岗位能力图谱 / 适岗度评价要求”五个页签，并支持返回建设中心首页。",
            ),
            (
                "JB-08",
                "基本信息与课程关联",
                "基本信息页展示岗位名称、职业、编码、层级、产业链-产业节点、岗位群、薪资、需求量、学历、经验等字段；“增加课程”用于维护岗位关联课程，课程关系需同步影响首页产业岗位课程图谱展示。",
            ),
            (
                "JB-09",
                "典型工作任务维护",
                "典型工作任务页提供“AI生成”和“手动添加”入口。手动添加弹窗至少包含任务名称、任务描述、关联能力项；任务名称和任务描述未填写时保存按钮禁用。",
            ),
            (
                "JB-10",
                "岗位能力图谱",
                "图谱页以岗位为中心，按任务驱动高亮知识/技能/素养三类能力项；切换任务时，相关能力高亮、无关能力置灰，保持用户可理解的能力映射关系。",
            ),
        ],
        state_rows=[
            ("空状态", "当前专业尚未初始化任何岗位数据", "首页展示“暂无岗位建设数据”、添加岗位按钮、模版导入按钮和三步引导，不可直接展示空白图谱。"),
            ("已添加岗位", "候选岗位已存在于建设中心", "候选卡片禁用并显示“已添加”，用户不能重复加入。"),
            ("确认禁用", "未勾选任何候选岗位", "确认按钮禁用；页面通过“已选 0 个”明确说明当前不可提交。"),
            ("搜索无结果", "添加岗位搜索后无匹配候选", "候选列表区域展示空结果提示，同时保留搜索词和“取消/关闭”入口。"),
            ("任务保存校验", "任务名称或任务描述为空", "保存按钮保持禁用，不允许提交空任务。"),
            ("课程无关联", "岗位当前未绑定课程", "基本信息页课程区域展示“暂无关联课程”，并保留“增加课程”按钮。"),
            ("操作失败", "添加岗位、删除岗位、保存任务失败", "保留当前弹窗或页面状态，提示失败原因并允许重试；用户已输入内容不丢失。"),
        ],
        flows=[
            "用户进入岗位建设中心，先判断当前专业是否已有岗位建设数据。",
            "若无数据，用户可通过“添加岗位”或“模版导入”完成初始化；若有数据，用户直接浏览图谱和岗位列表。",
            "用户在添加岗位弹窗中筛选候选岗位并确认导入，系统刷新岗位总数、图谱和列表。",
            "用户点击岗位卡片进入岗位详情页，补充课程关系、任务和能力项，并在图谱页查看任务与能力映射。",
            "建设结果沉淀完成后，成果页岗位中心直接复用相关岗位数据进行展示。",
        ],
        exclusions=[
            "一期不做完整审批流和版本对比。",
            "一期不做复杂批量导入模板配置后台。",
            "一期保留“AI建岗 / AI补全 / AI生成”入口，但不要求形成完整自动化闭环。",
        ],
        screenshots=[
            (SCREENSHOTS["job_build_empty"], "图 3  真实 demo 截图：岗位建设中心空状态", 6.25),
            (SCREENSHOTS["job_build_list"], "图 4  真实 demo 截图：岗位建设中心岗位列表区", 6.25),
        ],
    )

    add_module(
        doc=doc,
        title="七、模块需求四：查看成果页-岗位中心",
        purpose="该模块用于把岗位建设结果转化为可汇报、可浏览、可验收的成果页内容，强调展示效果和理解效率，而不是再次提供后台维护能力。",
        entries=[
            ("一级入口", "岗位中心顶部导航 -> 建设成果展示 -> 查看成果页"),
            ("栏目入口", "成果页顶部导航 -> 岗位中心"),
            ("页面对象", "关联岗位卡片轮播、洞察条、KPI、岗位建设路径、岗位产业图谱"),
        ],
        req_rows=[
            (
                "RJ-01",
                "成果页进入方式",
                "点击“建设成果展示”后需提供“查看成果页”入口；进入成果页后默认展示首页，点击顶部“岗位中心”栏目切换到岗位建设成果页面。",
            ),
            (
                "RJ-02",
                "关联岗位卡片轮播",
                "岗位中心页需展示岗位卡片轮播，包含上一个/下一个按钮、当前序号、页点切换。每张卡片至少展示岗位群、岗位名称、摘要、产业链、产业节点、所属职业、典型任务数、能力项数、关联课程数和岗课匹配度。",
            ),
            (
                "RJ-03",
                "卡片内图谱入口",
                "每张岗位卡片需提供“查看岗位能力图谱”按钮。点击后在当前页面下方图谱区切换为该岗位的能力图谱视图，并保留返回总图谱入口。",
            ),
            (
                "RJ-04",
                "岗位洞察与KPI",
                "岗位中心页需展示岗位洞察条和 KPI 信息，至少覆盖关联产业链、朝阳产业、开设趋势、就业规模等内容，用于支撑汇报解读。",
            ),
            (
                "RJ-05",
                "岗位建设路径",
                "成果页需用步骤化方式说明岗位建设路径，至少包括产业定位、岗位群聚焦、任务能力拆解和课程资源反哺等步骤，帮助非系统用户理解建设逻辑。",
            ),
            (
                "RJ-06",
                "岗位产业图谱",
                "成果页图谱区需展示产业链、产业节点、岗位群/岗位和课程的数量统计与图谱关系；当用户从岗位卡片切换到单岗位能力图谱后，需支持一键回到总图谱。",
            ),
            (
                "RJ-07",
                "成果页定位",
                "本页面仅用于展示，不承担岗位增删改、课程维护或配置发布职能；展示口径需与岗位建设中心保持一致。",
            ),
        ],
        state_rows=[
            ("无岗位成果", "成果页岗位中心无可展示岗位", "显示“暂无岗位建设成果”，保留页面标题、栏目导航和返回动作，不直接展示空白轮播。"),
            ("轮播边界", "点击上一张/下一张超出边界", "按循环切换处理，避免出现无内容页。"),
            ("单岗位图谱返回", "用户已进入单岗位能力图谱", "需保留“返回总图谱”动作，返回后恢复产业总图谱视图。"),
            ("其他栏目", "切换到非岗位中心栏目", "一期可展示空白占位栏目，不要求补充完整内容，但导航应可正常切换。"),
            ("加载失败", "成果数据拉取失败", "显示失败提示和重试入口；顶部成果页导航和岗位中心栏目结构不消失。"),
        ],
        flows=[
            "用户点击“建设成果展示”，选择“查看成果页”。",
            "进入成果页后，用户切换到顶部“岗位中心”栏目。",
            "用户通过轮播查看不同岗位成果卡片，并结合洞察条、KPI 和岗位建设路径理解建设成果。",
            "用户点击“查看岗位能力图谱”进入单岗位图谱，再返回总图谱继续浏览。",
        ],
        exclusions=[
            "一期不在成果页中直接编辑岗位或课程。",
            "一期不做成果发布审批、版本管理和外链权限控制。",
            "一期不要求补齐岗位中心以外所有成果栏目内容。",
        ],
        screenshots=[
            (SCREENSHOTS["results_job_center"], "图 5  真实 demo 截图：查看成果页-岗位中心", 6.25),
        ],
    )

    add_heading(doc, "八、一期验收关注点", 1)
    add_bullets(
        doc,
        [
            "四个模块入口清晰，能够形成“产业布局 -> 岗位分析 -> 岗位建设中心 -> 成果页岗位中心”的主链路。",
            "岗位建设中心需同时覆盖空状态和有数据状态，且添加岗位、删除岗位、任务维护等关键动作具备明确反馈。",
            "岗位分析和产业布局需具备基础筛选/查看能力，不能只停留在静态展示截图层面。",
            "成果页岗位中心需具备可讲解性，非系统用户能够通过岗位卡片、KPI 和图谱快速理解成果内容。",
            "异常态、禁用态和空态需有明确页面处理，不允许直接留白或静默失败。",
        ],
    )

    add_heading(doc, "九、截图说明", 1)
    add_table(
        doc,
        ["截图编号", "来源说明"],
        [
            ("图1", "本地 demo 页面：岗位数据调研-产业布局-产业链图谱"),
            ("图2", "本地 demo 页面：岗位数据调研-岗位分析-岗位画像分析"),
            ("图3", "本地 demo 页面：岗位建设中心空状态"),
            ("图4", "本地 demo 页面：岗位建设中心岗位列表区"),
            ("图5", "本地 demo 页面：查看成果页-岗位中心"),
        ],
        [1.2, 5.9],
        header_fill="F2F6FC",
    )

    doc.save(OUT_PATH)
    return OUT_PATH


if __name__ == "__main__":
    path = build_doc()
    print(path)
