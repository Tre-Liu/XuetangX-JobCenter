from __future__ import annotations

from collections import defaultdict
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from openpyxl import load_workbook
from PIL import Image, ImageDraw, ImageFont


ROOT = Path("/Users/liuhongzhe/Documents/专业建设/major-construction-platform")
INPUT_XLSX = ROOT / "outputs" / "prd" / "一期功能清单_v1.xlsx"
OUT_DIR = ROOT / "outputs" / "prd"
ASSET_DIR = OUT_DIR / "prd_assets"
SCREENSHOT_DIR = OUT_DIR / "current-demo-screenshots"
OUT_PATH = OUT_DIR / "岗位中心一期产品PRD_需求评审稿.docx"

FONT_CJK = Path("/System/Library/Fonts/STHeiti Medium.ttc")
FONT_CJK_LIGHT = Path("/System/Library/Fonts/STHeiti Light.ttc")

BLUE = "2E74B5"
DARK_BLUE = "1F4D78"
INK = "17233A"
MUTED = "5B677A"
LIGHT_BLUE = "EAF2FF"
LIGHT_GRAY = "F4F6F9"
BORDER = "D9E2F3"
GREEN = "10A37F"
PURPLE = "6C4BD9"


def pil_color(value: str) -> str:
    if value.startswith("#"):
        return value
    return f"#{value}"


def prd_text(value) -> str:
    return str(value or "").replace("智能建造工程专业岗位建设示例数据", "人工智能专业岗位建设示例数据")


def load_function_rows() -> list[dict[str, str]]:
    wb = load_workbook(INPUT_XLSX, data_only=True)
    ws = wb["一期功能清单"]
    rows: list[dict[str, str]] = []
    last_module = ""
    last_primary = ""
    for row in ws.iter_rows(min_row=3, values_only=True):
        if not any(row):
            continue
        module, primary, secondary, desc, note = row[:5]
        if module:
            last_module = str(module)
        if primary:
            last_primary = str(primary)
        if not secondary:
            continue
        rows.append(
            {
                "module": last_module,
                "primary": last_primary,
                "secondary": prd_text(secondary),
                "description": prd_text(desc),
                "note": prd_text(note),
            }
        )
    return rows


def grouped(rows: list[dict[str, str]]) -> dict[str, dict[str, list[dict[str, str]]]]:
    groups: dict[str, dict[str, list[dict[str, str]]]] = defaultdict(lambda: defaultdict(list))
    for item in rows:
        groups[item["module"]][item["primary"]].append(item)
    return groups


def font(size: int, bold: bool = False, light: bool = False) -> ImageFont.FreeTypeFont:
    path = FONT_CJK_LIGHT if light and FONT_CJK_LIGHT.exists() else FONT_CJK
    return ImageFont.truetype(str(path), size=size)


def rounded(draw: ImageDraw.ImageDraw, box, radius, fill, outline=None, width=1):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def draw_centered_text(draw: ImageDraw.ImageDraw, box, text: str, fnt, fill=INK, spacing=4):
    fill = pil_color(fill)
    x1, y1, x2, y2 = box
    max_width = x2 - x1 - 28
    lines: list[str] = []
    for part in text.split("\n"):
        current = ""
        for ch in part:
            test = current + ch
            if draw.textlength(test, font=fnt) <= max_width:
                current = test
            else:
                if current:
                    lines.append(current)
                current = ch
        if current:
            lines.append(current)
    line_h = fnt.size + spacing
    total_h = line_h * len(lines) - spacing
    y = y1 + ((y2 - y1) - total_h) / 2
    for line in lines:
        w = draw.textlength(line, font=fnt)
        draw.text((x1 + ((x2 - x1) - w) / 2, y), line, font=fnt, fill=fill)
        y += line_h


def draw_flow_diagram(path: Path):
    path.parent.mkdir(parents=True, exist_ok=True)
    img = Image.new("RGB", (1600, 640), "white")
    draw = ImageDraw.Draw(img)
    title_font = font(42, bold=True)
    label_font = font(25, bold=True)
    small_font = font(21, light=True)
    draw.text((60, 42), "一期业务闭环流程", font=title_font, fill=pil_color(INK))
    draw.text((60, 96), "从 CMS 数据初始化，到课程模型关联岗位能力项，形成需求评审范围内的最小闭环。", font=small_font, fill=pil_color(MUTED))

    boxes = [
        ("CMS数据初始化", "产业调研开通\n产业链选择"),
        ("产业布局", "产业链 / 区域\n政策 / 企业"),
        ("岗位分析", "画像 / 趋势\n新技术预判"),
        ("岗位建设中心", "图谱 / 岗位列表\n演示数据导入"),
        ("岗位详情维护", "基础信息 / 任务\n能力项 / 教务课"),
        ("AI课关联能力项", "AI课知识点\n关联岗位能力项"),
    ]
    colors = ["#F7FBFF", "#EAF2FF", "#EFFFF7", "#F3F0FF", "#FFF7E6", "#EEF7FF"]
    x = 45
    y = 230
    w = 225
    h = 160
    gap = 35
    for idx, (name, sub) in enumerate(boxes):
        bx = x + idx * (w + gap)
        rounded(draw, (bx, y, bx + w, y + h), 24, colors[idx], "#B8CCE4", 3)
        draw_centered_text(draw, (bx + 10, y + 26, bx + w - 10, y + 78), name, label_font, INK)
        draw_centered_text(draw, (bx + 14, y + 86, bx + w - 14, y + h - 18), sub, small_font, MUTED)
        if idx < len(boxes) - 1:
            sx = bx + w + 10
            ex = bx + w + gap - 10
            yy = y + h / 2
            draw.line((sx, yy, ex, yy), fill="#4D7CFE", width=5)
            draw.polygon([(ex, yy), (ex - 18, yy - 10), (ex - 18, yy + 10)], fill="#4D7CFE")

    rounded(draw, (210, 475, 1390, 570), 20, "#F4F6F9", "#D9E2F3", 2)
    draw.text((245, 501), "评审口径：CMS 只写数据初始化和产业链选择，不纳入报告生成、审批流、权限流、真实爬取、接口设计。", font=small_font, fill=pil_color(INK))
    img.save(path)


def draw_ia_diagram(path: Path, groups: dict[str, dict[str, list[dict[str, str]]]]):
    path.parent.mkdir(parents=True, exist_ok=True)
    modules = list(groups.keys())
    img = Image.new("RGB", (1600, 920), "white")
    draw = ImageDraw.Draw(img)
    title_font = font(42, bold=True)
    module_font = font(26, bold=True)
    item_font = font(20, light=True)
    draw.text((60, 42), "一期产品信息架构", font=title_font, fill=pil_color(INK))
    draw.text((60, 96), "以左侧导航模块为主线，展开到一级功能和关键页面能力。", font=item_font, fill=pil_color(MUTED))

    card_w = 460
    card_h = 235
    positions = [(70, 180), (570, 180), (1070, 180), (70, 490), (570, 490), (1070, 490)]
    palette = ["#F7FBFF", "#EAF2FF", "#EFFFF7", "#F3F0FF", "#FFF7E6", "#EEF7FF"]
    for idx, module in enumerate(modules[:6]):
        x, y = positions[idx]
        rounded(draw, (x, y, x + card_w, y + card_h), 22, palette[idx], "#B8CCE4", 2)
        draw.text((x + 28, y + 24), module, font=module_font, fill=pil_color(DARK_BLUE))
        primaries = list(groups[module].keys())[:7]
        yy = y + 78
        for primary in primaries:
            display_primary = "AI课岗位能力关联" if primary == "课程模型岗位能力" else primary
            draw.ellipse((x + 30, yy + 8, x + 42, yy + 20), fill="#4D7CFE")
            draw.text((x + 54, yy), display_primary, font=item_font, fill=pil_color(INK))
            yy += 32
    img.save(path)


def set_east_asia(run, name="Microsoft YaHei"):
    r_pr = run._element.get_or_add_rPr()
    fonts = r_pr.rFonts
    if fonts is None:
        fonts = OxmlElement("w:rFonts")
        r_pr.append(fonts)
    fonts.set(qn("w:eastAsia"), name)
    fonts.set(qn("w:ascii"), "Calibri")
    fonts.set(qn("w:hAnsi"), "Calibri")


def run_font(run, size=11, bold=False, color=None):
    run.font.name = "Calibri"
    run.font.size = Pt(size)
    run.bold = bold
    set_east_asia(run)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def paragraph_border_bottom(paragraph, color="2E74B5", size="8"):
    p_pr = paragraph._p.get_or_add_pPr()
    borders = p_pr.first_child_found_in("w:pBdr")
    if borders is None:
        borders = OxmlElement("w:pBdr")
        p_pr.append(borders)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), color)
    borders.append(bottom)


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def border_cell(cell, color=BORDER):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        el = borders.find(qn(f"w:{edge}"))
        if el is None:
            el = OxmlElement(f"w:{edge}")
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "6")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def set_cell_margins(cell, top=80, bottom=80, start=120, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("bottom", bottom), ("start", start), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_width(table, widths: list[int], indent=120):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), str(indent))
    tbl_ind.set(qn("w:type"), "dxa")

    grid = tbl.tblGrid
    if grid is None:
        grid = OxmlElement("w:tblGrid")
        tbl.insert(1, grid)
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")


def add_para(doc, text="", size=11, bold=False, color=None, before=0, after=6, align=None):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.10
    if text:
        r = p.add_run(text)
        run_font(r, size=size, bold=bold, color=color)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.space_before = Pt({1: 16, 2: 12, 3: 8}[level])
    p.paragraph_format.space_after = Pt({1: 8, 2: 6, 3: 4}[level])
    r = p.add_run(text)
    run_font(r, size={1: 16, 2: 13, 3: 12}[level], bold=True, color={1: BLUE, 2: BLUE, 3: DARK_BLUE}[level])
    return p


def add_bullets(doc, items: list[str]):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.line_spacing = 1.167
        r = p.add_run(item)
        run_font(r, size=11)


def add_table(doc, headers: list[str], rows: list[list[str]], widths: list[int], header_fill=LIGHT_GRAY):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_width(table, widths)
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade_cell(cell, header_fill)
        border_cell(cell)
        set_cell_margins(cell)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(header)
        run_font(r, size=10, bold=True, color=INK)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cell = cells[i]
            border_cell(cell)
            set_cell_margins(cell)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = cell.paragraphs[0]
            if i == 0 and len(val) <= 8:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(val)
            run_font(r, size=9 if len(val) > 80 else 10)
    set_table_width(table, widths)
    add_para(doc, after=2)
    return table


def add_callout(doc, title: str, text: str, fill="F4F6F9"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_width(table, [9360])
    cell = table.rows[0].cells[0]
    shade_cell(cell, fill)
    border_cell(cell, "C9D7EF")
    set_cell_margins(cell, top=140, bottom=140, start=180, end=180)
    p = cell.paragraphs[0]
    r = p.add_run(title)
    run_font(r, size=11, bold=True, color=DARK_BLUE)
    p.add_run("\n")
    r2 = p.add_run(text)
    run_font(r2, size=10, color=INK)
    add_para(doc, after=4)


def clean_text(value: str) -> str:
    return value.replace("None", "").strip()


def product_implementation(item: dict[str, str]) -> str:
    secondary = item["secondary"]
    primary = item["primary"]
    desc = clean_text(item["description"])
    if "弹窗" in secondary or "弹窗" in desc:
        prefix = "以弹窗承载该功能，用户在当前页面完成查看、选择或编辑后回到原页面，不跳出业务上下文。"
    elif any(word in secondary + desc for word in ["图谱", "桑基", "趋势", "矩阵", "KPI", "云"]):
        prefix = "以可视化区域承载该功能，页面默认展示核心指标、图形关系和说明信息，便于评审时直接观察数据含义。"
    elif any(word in secondary + desc for word in ["列表", "表格", "卡片", "明细"]):
        prefix = "以列表、表格或卡片承载该功能，统一展示对象名称、关键字段和可操作入口。"
    elif any(word in secondary + desc for word in ["搜索", "筛选"]):
        prefix = "以页面筛选工具承载该功能，用户输入条件后在当前列表或图谱范围内查看结果。"
    elif any(word in secondary + desc for word in ["导入", "下载", "模版"]):
        prefix = "以明确按钮入口承载该功能，入口文案与当前 demo 保持一致，避免误解为通用后台能力。"
    elif any(word in secondary + desc for word in ["编辑", "添加", "删除", "保存"]):
        prefix = "以表单和操作按钮承载该功能，用户在当前业务对象范围内完成内容维护。"
    else:
        prefix = f"在“{primary}”中作为独立展示或操作能力落地，保持与当前 demo 的页面层级一致。"
    business_note = business_course_context(item)
    return f"{prefix}具体内容：{desc}{business_note}"


def business_course_context(item: dict[str, str]) -> str:
    module = item["module"]
    secondary = item["secondary"]
    primary = item["primary"]
    if secondary in {"关联课程展示", "增加关联课程弹窗", "删除关联课程"}:
        return " 业务口径：此处课程为教务课，来源于学校教务课程库，用于维护岗位与正式培养方案课程之间的关系。"
    if module == "课程关联岗位能力项" or "课程模型" in primary or "课程模型" in item["description"]:
        return " 业务口径：此处课程为 AI 课，对应课程管理中的“关联的AI课”对象，用于 AI 课知识点与岗位能力项建立关联，不与岗位详情中的教务课混用。"
    return ""


def interaction_detail(item: dict[str, str]) -> str:
    text = f"{item['secondary']}。{clean_text(item['description'])}。{clean_text(item['note'])}"
    details: list[str] = []
    patterns = [
        ("点击", "点击入口后在当前页面、弹窗或独立页面展示对应内容。"),
        ("悬停", "悬停时高亮相关节点或连线，并展示说明信息。"),
        ("搜索", "输入关键词后触发筛选，结果范围限定在当前页面数据。"),
        ("筛选", "选择筛选条件后刷新当前列表或图表展示。"),
        ("勾选", "勾选或取消勾选后同步更新已选数量和按钮状态。"),
        ("选择", "选择对象后进入已选态，支持取消或切换。"),
        ("切换", "切换任务、岗位或视图后，右侧内容、连线和清单同步变化。"),
        ("分页", "支持上一页、下一页、页码和总页数展示。"),
        ("禁用", "不满足操作条件时按钮置灰禁用，并避免产生无效提交。"),
        ("关闭", "关闭弹窗时回到原页面，并清理临时输入或选择状态。"),
        ("保存", "保存前进行必要字段校验，保存后刷新当前页面展示。"),
        ("删除", "删除前给出确认或明确操作入口，删除后当前页面状态同步更新。"),
        ("导入", "导入成功后将示例或解析内容写入当前业务对象并刷新展示。"),
        ("下载", "点击下载后获得当前 demo 指定模板文件。"),
        ("解析", "解析失败时展示问题清单，解析通过后才允许继续导入。"),
        ("输入", "输入框需要保留示例或占位提示，必要时限制输入格式。"),
        ("返回", "返回入口应回到上一级图谱或列表，保留业务链路可理解。"),
    ]
    for keyword, sentence in patterns:
        if keyword in text and sentence not in details:
            details.append(sentence)
    if not details:
        details.append("进入页面后默认展示，不要求额外操作；如页面存在卡片或字段，按当前 demo 状态展示。")
    return " ".join(details[:4])


def data_and_rules(item: dict[str, str]) -> str:
    note = clean_text(item["note"])
    course_context = business_course_context(item).strip()
    if note.startswith("数据维度："):
        return f"{note} {course_context}".strip()
    if "字段" in note or "示例" in note or "来源" in note or "按钮" in note or "禁用" in note or "校验" in note:
        return f"{note} {course_context}".strip()
    desc = clean_text(item["description"])
    if any(word in desc for word in ["展示", "显示"]):
        return f"数据与字段以当前 demo 展示内容为准，评审时重点确认字段完整性和含义清晰度。 {course_context}".strip()
    return f"按当前 demo 的页面状态和业务对象范围实现，不额外扩展后台配置项。 {course_context}".strip()


def acceptance_criteria(item: dict[str, str]) -> str:
    secondary = item["secondary"]
    desc = clean_text(item["description"])
    note = clean_text(item["note"])
    checks = [f"评审时应能在对应页面定位到“{secondary}”。"]
    if any(word in secondary + desc for word in ["搜索", "筛选"]):
        checks.append("输入或选择条件后，当前列表/图表展示范围应发生可见变化或保持可解释的空结果。")
    elif any(word in secondary + desc for word in ["弹窗", "详情"]):
        checks.append("触发后应打开对应弹窗或详情区域，字段信息完整，关闭后返回原页面。")
    elif any(word in secondary + desc for word in ["图谱", "矩阵", "连线"]):
        checks.append("图谱节点、分类、连线和高亮关系应能被清楚识别。")
    elif any(word in secondary + desc for word in ["导入", "下载", "解析"]):
        checks.append("入口、成功态、失败提示或禁用规则应与当前 demo 行为一致。")
    elif any(word in secondary + desc for word in ["编辑", "添加", "删除", "保存"]):
        checks.append("新增、编辑、删除或保存后的页面状态应立即体现变化。")
    if note and not note.startswith("数据维度："):
        checks.append(f"注意边界：{note}")
    return " ".join(checks[:3])


def add_image(doc, path: Path, caption: str, width=6.25):
    if not path.exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(str(path), width=Inches(width))
    cap = add_para(doc, caption, size=9, color=MUTED, after=8, align=WD_ALIGN_PARAGRAPH.CENTER)
    cap.runs[0].italic = True


MODULE_SCREENSHOTS = {
    "CMS数据初始化": [
        ("00-cms-uninitialized-prompt.png", "岗位中心：产业调研数据未初始化提示与前往 CMS 初始化入口"),
        ("22-cms-init-empty.png", "CMS：产业调研管理待初始化状态"),
        ("23-cms-chain-recommendations.png", "CMS：数据初始化后的推荐产业链列表"),
        ("24-cms-chain-selected.png", "CMS：选择产业链后回写初始化状态"),
    ],
    "产业布局": [
        ("01-industry-chain.png", "产业布局：产业链图谱与产业链结构分析"),
        ("02-industry-region.png", "产业布局：区域产业分析与区域热度分布"),
        ("03-industry-policy.png", "产业布局：产业政策库与政策影响分析"),
        ("04-industry-company.png", "产业布局：产业企业库与企业详情入口"),
    ],
    "岗位分析": [
        ("05-job-portrait.png", "岗位分析：岗位画像搜索、画像洞察与岗位列表"),
        ("06-job-portrait-detail.png", "岗位分析：岗位详情弹窗与三维能力分析"),
        ("07-job-demand.png", "岗位分析：招聘需求趋势与数据维度看板"),
        ("08-job-forecast.png", "岗位分析：新岗位新技术预判与人才培养方向建议"),
        ("09-job-competency-map.png", "岗位分析：独立岗位能力图谱与任务-能力项联动"),
    ],
    "岗位建设中心": [
        ("10-build-empty.png", "岗位建设中心：未导入岗位时的空状态与建设引导"),
        ("11-build-add-job-dialog.png", "岗位建设中心：添加岗位弹窗与候选岗位选择"),
        ("12-build-after-add.png", "岗位建设中心：导入岗位后的产业岗位课程图谱"),
    ],
    "岗位详情维护": [
        ("13-job-detail-basic.png", "岗位详情维护：岗位基本信息与关联教务课区域"),
        ("14-job-detail-edu-course-dialog.png", "岗位详情维护：增加关联课程弹窗，此处课程为教务课"),
        ("15-job-detail-tasks.png", "岗位详情维护：典型工作任务列表与任务维护"),
        ("16-job-detail-abilities.png", "岗位详情维护：岗位能力项列表与能力项维护"),
        ("17-job-detail-ability-import.png", "岗位详情维护：岗位能力项导入弹窗"),
        ("18-job-detail-ability-map.png", "岗位详情维护：岗位能力图谱与任务-能力项联动"),
    ],
    "课程关联岗位能力项": [
        ("19-course-model.png", "课程模型：课程知识图谱与知识点编辑入口"),
        ("20-course-ai-ability-dialog.png", "课程模型：关联岗位能力弹窗，此处课程对象为 AI 课"),
    ],
}


def resolve_screenshot(relative_path: str) -> Path:
    return (SCREENSHOT_DIR / relative_path).resolve()


def add_module_screenshots(doc: Document, module: str):
    screenshots = MODULE_SCREENSHOTS.get(module, [])
    available = [(resolve_screenshot(path), caption) for path, caption in screenshots if resolve_screenshot(path).exists()]
    if not available:
        return
    add_heading(doc, f"{module}Demo 截图", 3)
    add_para(doc, "以下截图用于辅助评审页面结构、主要交互和业务对象，具体需求口径仍以本节功能说明为准。", size=10, color=MUTED, after=4)
    for path, caption in available:
        add_image(doc, path, caption, width=6.35)


def setup_doc(doc: Document):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK)
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for name, size, color in [
        ("Heading 1", 16, BLUE),
        ("Heading 2", 13, BLUE),
        ("Heading 3", 12, DARK_BLUE),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    header = section.header.paragraphs[0]
    header.text = ""
    header.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = header.add_run("岗位中心一期产品 PRD | 需求评审稿")
    run_font(r, size=9, color=MUTED)
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = footer.add_run("内部评审资料")
    run_font(r, size=9, color=MUTED)


def add_cover(doc: Document):
    add_para(doc, "产品需求文档（PRD）", size=12, bold=True, color=BLUE, before=16, after=8)
    title = add_para(doc, "岗位中心一期产品需求说明", size=24, bold=True, color=INK, before=4, after=4)
    paragraph_border_bottom(title, BLUE, "10")
    add_para(doc, "需求评审稿 | 基于《一期功能清单_v1》修订", size=12, color=MUTED, after=18)
    meta = [
        ["产品阶段", "一期需求评审"],
        ["产品范围", "CMS数据初始化、产业布局、岗位分析、岗位建设中心、岗位详情维护、课程关联岗位能力项"],
        ["示例专业", "一期先以“人工智能专业”为例进行数据准备"],
        ["不纳入范围", "产业调研报告生成、审批/权限流程、真实爬取、后台接口设计、全自动分析"],
        ["文档日期", str(date.today())],
        ["使用对象", "产品、研发、测试、实施、业务评审人员"],
    ]
    add_table(doc, ["字段", "内容"], meta, [1600, 7600], header_fill="EAF2FF")
    add_callout(
        doc,
        "评审摘要",
        "本 PRD 只描述当前 demo 已呈现并纳入一期的产品能力，重点用于确认页面范围、业务闭环、关键交互、数据展示维度和验收口径。文档刻意不展开技术实现、接口协议、库表结构和权限策略。",
        fill="F7FBFF",
    )


def add_overview(doc: Document, groups: dict[str, dict[str, list[dict[str, str]]]], flow_path: Path, ia_path: Path):
    add_heading(doc, "1. 产品背景与目标", 1)
    add_para(
        doc,
        "岗位中心一期围绕专业岗位建设场景展开，目标是把 CMS 数据初始化、产业与岗位侧数据、岗位画像分析、岗位建设图谱、岗位详情维护以及课程知识点关联岗位能力项串成一条可评审、可演示、可验收的产品闭环。",
    )
    add_bullets(
        doc,
        [
            "让评审人员能看清岗位数据从哪里来、如何被分析、如何沉淀为岗位建设资产。",
            "让专业负责人能围绕岗位、任务、能力项和课程支撑关系进行讨论和修订。",
            "让研发和测试能按页面、交互、数据维度和边界条件拆分一期交付内容。",
        ]
    )

    add_heading(doc, "2. 一期范围", 1)
    module_rows = []
    for module, primaries in groups.items():
        second_count = sum(len(items) for items in primaries.values())
        module_rows.append([module, "、".join(primaries.keys()), f"{second_count} 项"])
    add_table(doc, ["模块", "一级功能", "功能项数量"], module_rows, [1700, 6000, 1660])
    add_callout(
        doc,
        "范围边界",
        "一期不做产业调研报告生成。CMS 仅纳入产业调研数据初始化、推荐产业链选择和初始化状态回写，不纳入审批/权限、真实爬取、后台接口设计和全自动分析。本期聚焦岗位中心主链路。",
        fill="FFF7E6",
    )
    add_callout(
        doc,
        "数据准备口径",
        "本期先以“人工智能专业”为示例专业进行数据准备，围绕该专业补齐 CMS 产业调研初始化状态、产业链/产业节点、岗位群/岗位、典型工作任务、岗位能力项、教务课、AI课及课程知识点关联岗位能力项等演示数据。当前 demo 截图用于说明页面结构与交互方式，正式评审数据口径以人工智能专业为准。",
        fill="F7FBFF",
    )
    add_callout(
        doc,
        "课程业务口径",
        "岗位详情维护中的“增加关联课程”使用教务课，代表学校教务课程库中的正式课程；课程管理列表中的“关联的AI课”使用 AI 课，代表课程模型侧可用于知识点分析和能力项关联的课程。两类课程都属于一期业务闭环，但在 PRD 中必须分开描述，避免把教务课和 AI 课混为同一个对象。",
        fill="F7FBFF",
    )

    add_heading(doc, "3. 业务闭环", 1)
    add_image(doc, flow_path, "图 1 一期业务闭环流程", width=6.4)
    add_para(
        doc,
        "闭环从 CMS 产业调研数据初始化开始，完成产业链选择后进入产业与岗位数据展示，通过岗位画像、招聘需求趋势和新岗位新技术预判形成分析依据，再进入岗位建设中心维护岗位、任务、能力项和课程关系，最后在课程模型中完成课程知识点与岗位能力项的关联。",
    )

    add_heading(doc, "4. 产品信息架构", 1)
    add_image(doc, ia_path, "图 2 一期产品信息架构", width=6.4)


def add_module_sections(doc: Document, groups: dict[str, dict[str, list[dict[str, str]]]]):
    add_heading(doc, "5. 模块需求说明", 1)
    module_purpose = {
        "CMS数据初始化": "用于开通当前专业产业调研数据，选择本期使用的产业链方向，是产业布局和岗位分析正常展示的前置条件。",
        "产业布局": "用于说明岗位建设所依托的产业链、区域、政策和企业背景，是岗位分析和岗位建设的数据展示入口。",
        "岗位分析": "用于完成岗位画像、招聘趋势、新岗位新技术三类分析，是一期最重要的数据可视化需求区域。",
        "岗位建设中心": "用于将岗位、产业节点、课程关系组织成建设图谱和岗位列表，并支持从演示数据或候选岗位进入建设。",
        "岗位详情维护": "用于维护单个岗位的基础信息、关联教务课、典型工作任务、岗位能力项、能力图谱和适岗度展示。",
        "课程关联岗位能力项": "用于在课程模型中把 AI 课知识点与岗位能力项建立关系，形成课程支撑岗位能力的闭环。",
    }
    module_impl = {
        "CMS数据初始化": [
            "岗位中心产业调研页面在未初始化时先展示空状态提示，引导用户前往 CMS。",
            "CMS 产业调研管理页提供数据初始化按钮，初始化后展示推荐产业链列表、匹配度、推荐原因和标签。",
            "用户选择至少 1 条产业链后回写初始化状态，返回岗位中心后产业布局和岗位分析页面展示对应数据。",
            "本模块只写当前 demo 的初始化状态、推荐列表、选择/取消、分页和状态回写，不写真实采集任务、审批流、接口设计或后台调度。",
        ],
        "产业布局": [
            "页面按左侧导航进入产业调研相关区域，需在 CMS 初始化完成后展示产业链图谱、区域产业分析、政策库和企业库。",
            "实现上以可视化和说明型内容为主，包含产业链图谱、区域产业分析、政策库和企业库。",
            "本模块不做数据采集任务、报告生成和后台配置，只承担岗位建设前置依据展示；未初始化时统一展示 CMS 初始化引导。",
        ],
        "岗位分析": [
            "页面按岗位画像、岗位能力图谱、招聘需求趋势、新岗位新技术预判组织。",
            "实现上以岗位搜索、岗位卡片、详情弹窗、独立图谱页和趋势图表串联，形成从岗位发现到能力拆解的分析链路。",
            "数据可视化必须写清维度，包括岗位、薪资、需求量、能力类别、任务、月份、技能热度、城市等。",
        ],
        "岗位建设中心": [
            "页面承接岗位分析结果，提供空状态引导、演示数据导入、候选岗位添加、产业岗位课程图谱和岗位列表。",
            "实现上以建设中心内状态管理为主，导入演示数据后展示岗位、任务、能力项和课程关系。",
            "AI 建岗只保留入口按钮，模版导入只指导入演示数据，不写通用文件上传和智能生成流程。",
        ],
        "岗位详情维护": [
            "页面围绕单个岗位维护，包含基本信息、关联教务课、典型工作任务、岗位能力项、岗位能力图谱和适岗度展示。",
            "增加关联课程弹窗中的课程为教务课，应展示教务课程名称、课程编码和已关联状态。",
            "实现上以表单弹窗、列表维护、文件导入、确认弹窗和图谱联动为主。",
            "本模块需要重点验收字段校验、按钮禁用、导入错误提示、能力项改名同步和删除引用同步。",
        ],
        "课程关联岗位能力项": [
            "页面入口来自课程模型知识点详情抽屉，目标是把 AI 课知识点与岗位能力项建立关联。",
            "课程管理列表中的“关联的AI课”列展示教务课与 AI 课的业务关联状态，PRD 中应明确此处课程对象为 AI 课。",
            "实现上以关联弹窗承载，左侧选岗位，右侧按知识、技能、素养勾选能力项。",
            "关闭、取消、确认关联等内部状态需要写清，避免临时勾选残留影响下一次操作。",
        ],
    }
    for module, primaries in groups.items():
        add_heading(doc, module, 2)
        add_para(doc, module_purpose.get(module, "该模块用于支撑一期岗位中心业务闭环。"))
        add_module_screenshots(doc, module)
        add_heading(doc, f"{module}产品实现方式", 3)
        add_bullets(doc, module_impl.get(module, ["按当前 demo 页面结构实现，不额外扩展未呈现能力。"]))

        primary_rows = []
        for primary, items in primaries.items():
            data_notes = []
            interactions = []
            for item in items:
                note = item["note"]
                desc = item["description"]
                if "数据维度" in note:
                    data_notes.append(note.replace("数据维度：", ""))
                if any(keyword in desc for keyword in ["点击", "悬停", "搜索", "勾选", "打开", "切换", "禁用", "关闭", "保存"]):
                    interactions.append(item["secondary"])
            primary_rows.append(
                [
                    primary,
                    "、".join([item["secondary"] for item in items[:6]]) + ("等" if len(items) > 6 else ""),
                    "；".join(data_notes[:2]) if data_notes else "以页面展示字段为准",
                    "、".join(interactions[:4]) if interactions else "以展示为主",
                ]
            )
        add_table(doc, ["一级功能", "页面/能力构成", "核心数据维度", "关键交互"], primary_rows, [1500, 3100, 2600, 2160])

        add_heading(doc, f"{module}逐功能说明", 3)
        for primary, items in primaries.items():
            add_para(doc, primary, size=11, bold=True, color=DARK_BLUE, before=6, after=4)
            detail_rows = [
                [
                    item["secondary"],
                    product_implementation(item),
                    interaction_detail(item),
                    f"{data_and_rules(item)}\n{acceptance_criteria(item)}",
                ]
                for item in items
            ]
            add_table(
                doc,
                ["功能点", "如何实现", "交互细节", "数据/验收口径"],
                detail_rows,
                [1500, 3200, 2300, 2360],
                header_fill="EAF2FF",
            )

        add_heading(doc, f"{module}评审要点", 3)
        review_points = []
        if module == "岗位分析":
            review_points = [
                "岗位画像分析、招聘需求趋势、新岗位新技术预判三块必须同时保留。",
                "每张可视化图或表格都需要明确数据维度，避免只写展示效果。",
                "岗位详情弹窗和独立岗位能力图谱属于一期内交互，不应遗漏。",
            ]
        elif module == "CMS数据初始化":
            review_points = [
                "CMS 初始化是产业调研页面展示数据的前置条件，需要验收未初始化提示、跳转、初始化、产业链选择和状态回写。",
                "产业链推荐仅按当前 demo 展示和选择能力验收，不写真实采集任务、后台调度和接口设计。",
                "初始化完成后，产业布局和岗位分析页面应从空状态切换为对应数据展示。",
            ]
        elif module == "岗位建设中心":
            review_points = [
                "模版导入仅指导入演示数据，不是通用 Excel 上传。",
                "AI 建岗当前只作为入口，不写完整生成流程。",
                "图谱的节点悬停、节点进入详情、返回产业图谱属于需要验收的交互。",
            ]
        elif module == "岗位详情维护":
            review_points = [
                "能力项导入、编辑、删除和改名同步任务引用是当前 demo 新增能力。",
                "关联课程影响岗位课程图谱，应作为产品行为写清。",
                "适岗度评价只做展示，不写评价算法和计算规则。",
            ]
        elif module == "课程关联岗位能力项":
            review_points = [
                "课程模型内的岗位能力项关联弹窗是一期闭环终点。",
                "需要评审岗位搜索、岗位切换、能力项勾选、确认关联和关闭清空草稿等交互。",
            ]
        else:
            review_points = [
                "该模块以 demo 展示数据为边界，不写实时采集和后台管理。",
                "筛选、搜索、悬停、详情等交互只按页面实际出现内容描述。",
            ]
        add_bullets(doc, review_points)


def add_review_and_appendix(doc: Document, rows: list[dict[str, str]]):
    add_heading(doc, "6. 需求评审检查清单", 1)
    checklist = [
        ["范围确认", "确认一期不包含产业调研报告生成、真实数据抓取、后台接口设计、审批流、权限流。CMS 仅包含数据初始化和产业链选择。"],
        ["页面确认", "确认 CMS数据初始化、产业布局、岗位分析、岗位建设中心、岗位详情维护、课程关联岗位能力项均在本期范围。"],
        ["数据确认", "确认每个可视化、卡片、表格的核心数据维度已写清。"],
        ["交互确认", "确认点击、搜索、悬停、切换、勾选、禁用、关闭等内部交互已被记录。"],
        ["验收确认", "确认当前 demo 入口型能力只按入口验收，不扩展生成、推荐、自动化流程。"],
    ]
    add_table(doc, ["检查项", "确认内容"], checklist, [1800, 7560])

    add_heading(doc, "7. 附录：功能清单摘要", 1)
    add_para(doc, "以下摘要自动取自《一期功能清单_v1.xlsx》，用于评审时快速核对模块和功能项。")
    compact_rows = [[r["module"], r["primary"], r["secondary"]] for r in rows]
    # Split the appendix table to avoid one huge table becoming hard to review.
    for index in range(0, len(compact_rows), 28):
        part = compact_rows[index:index + 28]
        add_table(doc, ["模块", "一级功能", "二级功能"], part, [1800, 2500, 5060])


def build():
    rows = load_function_rows()
    groups = grouped(rows)
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    flow_path = ASSET_DIR / "phase1_business_flow.png"
    ia_path = ASSET_DIR / "phase1_information_architecture.png"
    draw_flow_diagram(flow_path)
    draw_ia_diagram(ia_path, groups)

    doc = Document()
    setup_doc(doc)
    add_cover(doc)
    add_overview(doc, groups, flow_path, ia_path)
    add_module_sections(doc, groups)
    add_review_and_appendix(doc, rows)
    doc.save(OUT_PATH)
    print(OUT_PATH)


if __name__ == "__main__":
    build()
