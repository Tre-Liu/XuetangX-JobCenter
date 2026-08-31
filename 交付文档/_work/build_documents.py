from __future__ import annotations

from pathlib import Path
from typing import Iterable, Sequence

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/liuhongzhe/Desktop/学堂/专业建设/Codex工程/交付文档")
SHOT_DIR = ROOT / "系统截图"
WORK_DIR = ROOT / "_work"
WORK_DIR.mkdir(parents=True, exist_ok=True)

TECH_DOC = ROOT / "视界学伴关键技术说明.docx"
FUNC_DOC = ROOT / "视界学伴功能说明书.docx"

BODY_CN = "Arial Unicode MS"
HEAD_CN = "Arial Unicode MS"
LATIN = "Arial Unicode MS"
FONT_FILE = "/System/Library/Fonts/STHeiti Medium.ttc"

NAVY = "163F3A"
TEAL = "2C6C62"
GOLD = "C9973A"
INK = "24313A"
MUTED = "66747D"
LINE = "D8E1DE"
LIGHT = "F3F7F5"
PALE_GOLD = "FBF3E3"
WHITE = "FFFFFF"

PAGE_WIDTH_DXA = 12240
CONTENT_WIDTH_DXA = 9360
TABLE_INDENT_DXA = 120


def rgb(hex_color: str) -> RGBColor:
    return RGBColor.from_string(hex_color)


def set_run_font(run, *, cn: str = BODY_CN, latin: str = LATIN, size: float | None = None,
                 color: str | None = None, bold: bool | None = None, italic: bool | None = None):
    run.font.name = latin
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:eastAsia"), cn)
    rfonts.set(qn("w:cs"), latin)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    return run


def set_cell_shading(cell, fill: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top: int = 80, start: int = 120, bottom: int = 80, end: int = 120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    existing = tr_pr.find(qn("w:tblHeader"))
    if existing is not None:
        existing.set(qn("w:val"), "true")
        return
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def set_cell_width(cell, width_dxa: int):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa: Sequence[int], indent_dxa: int = TABLE_INDENT_DXA):
    assert sum(widths_dxa) == CONTENT_WIDTH_DXA, (widths_dxa, sum(widths_dxa))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    tbl_pr = table._tbl.tblPr
    for child_name in ("tblW", "tblInd", "tblLayout"):
        child = tbl_pr.find(qn(f"w:{child_name}"))
        if child is None:
            child = OxmlElement(f"w:{child_name}")
            tbl_pr.append(child)
        if child_name == "tblW":
            child.set(qn("w:w"), str(CONTENT_WIDTH_DXA))
            child.set(qn("w:type"), "dxa")
        elif child_name == "tblInd":
            child.set(qn("w:w"), str(indent_dxa))
            child.set(qn("w:type"), "dxa")
        else:
            child.set(qn("w:type"), "fixed")
    grid = table._tbl.tblGrid
    for node in list(grid):
        grid.remove(node)
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            set_cell_width(cell, widths_dxa[min(idx, len(widths_dxa) - 1)])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if table.rows:
        set_repeat_table_header(table.rows[0])


def set_table_borders(table, color: str = LINE, size: int = 6):
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), str(size))
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_cell_text(cell, text: str, *, bold: bool = False, color: str = INK,
                  size: float = 9.5, align=WD_ALIGN_PARAGRAPH.LEFT):
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.15
    p.clear()
    set_run_font(p.add_run(text), size=size, color=color, bold=bold)


def configure_styles(doc: Document):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = LATIN
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_CN)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10
    normal.paragraph_format.widow_control = True

    for name, size, color, before, after in (
        ("Heading 1", 16, TEAL, 16, 8),
        ("Heading 2", 13, TEAL, 12, 6),
        ("Heading 3", 11.5, NAVY, 8, 4),
    ):
        style = styles[name]
        style.font.name = LATIN
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), HEAD_CN)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    title = styles["Title"]
    title.font.name = LATIN
    title.font.size = Pt(28)
    title.font.bold = True
    title.font.color.rgb = rgb(NAVY)
    title._element.rPr.rFonts.set(qn("w:eastAsia"), HEAD_CN)
    title.paragraph_format.space_after = Pt(8)

    subtitle = styles["Subtitle"]
    subtitle.font.name = LATIN
    subtitle.font.size = Pt(13)
    subtitle.font.color.rgb = rgb(MUTED)
    subtitle._element.rPr.rFonts.set(qn("w:eastAsia"), HEAD_CN)
    subtitle.paragraph_format.space_after = Pt(12)

    for list_name in ("List Bullet", "List Number"):
        style = styles[list_name]
        style.font.name = LATIN
        style.font.size = Pt(10.5)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_CN)
        style.paragraph_format.left_indent = Inches(0.5)
        style.paragraph_format.first_line_indent = Inches(-0.25)
        style.paragraph_format.space_after = Pt(5)
        style.paragraph_format.line_spacing = 1.15

    if "Figure Caption" not in styles:
        cap = styles.add_style("Figure Caption", WD_STYLE_TYPE.PARAGRAPH)
    else:
        cap = styles["Figure Caption"]
    cap.font.name = LATIN
    cap.font.size = Pt(9)
    cap.font.color.rgb = rgb(MUTED)
    cap._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_CN)
    cap.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.paragraph_format.space_before = Pt(4)
    cap.paragraph_format.space_after = Pt(10)
    cap.paragraph_format.keep_with_next = False

    if "Source Note" not in styles:
        source = styles.add_style("Source Note", WD_STYLE_TYPE.PARAGRAPH)
    else:
        source = styles["Source Note"]
    source.font.name = LATIN
    source.font.size = Pt(8.5)
    source.font.color.rgb = rgb(MUTED)
    source._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_CN)
    source.paragraph_format.space_before = Pt(4)
    source.paragraph_format.space_after = Pt(4)


def add_field(paragraph, field_code: str, display: str = "1"):
    run = paragraph.add_run()
    fld_char = OxmlElement("w:fldChar")
    fld_char.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = field_code
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = display
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([fld_char, instr, separate, text, end])
    set_run_font(run, size=8.5, color=MUTED)


def configure_page(doc: Document, short_title: str, version: str = "V1.0"):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.right_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)
    section.different_first_page_header_footer = True

    header = section.header
    p = header.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(3)
    set_run_font(p.add_run(f"视界学伴  |  {short_title}  |  {version}"), size=8.5, color=MUTED)
    ppr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), LINE)
    pbdr.append(bottom)
    ppr.append(pbdr)

    first_header = section.first_page_header
    first_header.paragraphs[0].clear()

    footer = section.footer
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    fp.paragraph_format.space_before = Pt(2)
    set_run_font(fp.add_run("第 "), size=8.5, color=MUTED)
    add_field(fp, "PAGE")
    set_run_font(fp.add_run(" 页"), size=8.5, color=MUTED)

    first_footer = section.first_page_footer
    first_footer.paragraphs[0].clear()


def add_paragraph(doc: Document, text: str, *, bold_prefix: str | None = None,
                  align=WD_ALIGN_PARAGRAPH.LEFT, after: float = 6, line: float = 1.10):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    if bold_prefix and text.startswith(bold_prefix):
        set_run_font(p.add_run(bold_prefix), bold=True, color=INK)
        set_run_font(p.add_run(text[len(bold_prefix):]), color=INK)
    else:
        set_run_font(p.add_run(text), color=INK)
    return p


def add_bullets(doc: Document, items: Iterable[str]):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        set_run_font(p.add_run(item), color=INK)


def new_decimal_num_id(doc: Document) -> int:
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(node.get(qn("w:abstractNumId"))) for node in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(node.get(qn("w:numId"))) for node in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    level = OxmlElement("w:lvl")
    level.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    level.append(start)
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    level.append(num_fmt)
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1.")
    level.append(lvl_text)
    lvl_jc = OxmlElement("w:lvlJc")
    lvl_jc.set(qn("w:val"), "left")
    level.append(lvl_jc)
    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "720")
    tabs.append(tab)
    ppr.append(tabs)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "720")
    ind.set(qn("w:hanging"), "360")
    ppr.append(ind)
    level.append(ppr)
    abstract.append(level)
    numbering.append(abstract)

    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)
    return num_id


def add_numbers(doc: Document, items: Iterable[str]):
    num_id = new_decimal_num_id(doc)
    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.25)
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing = 1.15
        ppr = p._p.get_or_add_pPr()
        num_pr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        num_id_node = OxmlElement("w:numId")
        num_id_node.set(qn("w:val"), str(num_id))
        num_pr.append(ilvl)
        num_pr.append(num_id_node)
        ppr.insert(0, num_pr)
        set_run_font(p.add_run(item), color=INK)


def add_callout(doc: Document, label: str, text: str, *, fill: str = LIGHT, accent: str = TEAL):
    table = doc.add_table(rows=1, cols=1)
    set_table_geometry(table, [CONTENT_WIDTH_DXA], indent_dxa=180)
    set_table_borders(table, color=fill, size=0)
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, top=150, start=180, bottom=150, end=180)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    set_run_font(p.add_run(label), cn=HEAD_CN, size=9.5, bold=True, color=accent)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_before = Pt(0)
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.15
    set_run_font(p2.add_run(text), size=10, color=INK)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_table(doc: Document, headers: Sequence[str], rows: Sequence[Sequence[str]], widths: Sequence[int]):
    table = doc.add_table(rows=1, cols=len(headers))
    set_table_geometry(table, widths)
    set_table_borders(table)
    set_repeat_table_header(table.rows[0])
    for idx, header in enumerate(headers):
        set_cell_shading(table.rows[0].cells[idx], LIGHT)
        set_cell_text(table.rows[0].cells[idx], header, bold=True, color=NAVY, size=9.3,
                      align=WD_ALIGN_PARAGRAPH.CENTER)
    for row in rows:
        cells = table.add_row().cells
        for idx, text in enumerate(row):
            set_cell_text(cells[idx], str(text), size=9.1,
                          align=WD_ALIGN_PARAGRAPH.CENTER if idx == 0 else WD_ALIGN_PARAGRAPH.LEFT)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_picture(doc: Document, image_path: Path, caption: str, alt_text: str,
                width: float = 6.35, page_break_before: bool = False):
    if page_break_before:
        doc.add_page_break()
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.keep_with_next = True
    run = p.add_run()
    inline_shape = run.add_picture(str(image_path), width=Inches(width))
    doc_pr = inline_shape._inline.docPr
    doc_pr.set("descr", alt_text)
    cap = doc.add_paragraph(style="Figure Caption")
    set_run_font(cap.add_run(caption), size=9, color=MUTED)


def add_cover(doc: Document, title: str, subtitle: str, doc_type: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(92)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run("VISION CAREER DEVELOPMENT SYSTEM"), cn=HEAD_CN, size=9.5,
                 color=GOLD, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    set_run_font(p.add_run("视界学伴"), cn=HEAD_CN, size=15, color=TEAL, bold=True)

    p = doc.add_paragraph(style="Title")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run(title), cn=HEAD_CN, size=28, color=NAVY, bold=True)

    p = doc.add_paragraph(style="Subtitle")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_run_font(p.add_run(subtitle), cn=HEAD_CN, size=13, color=MUTED)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(96)
    set_run_font(p.add_run(doc_type), cn=HEAD_CN, size=10.5, color=GOLD, bold=True)

    meta = doc.add_table(rows=3, cols=2)
    set_table_geometry(meta, [4680, 4680], indent_dxa=120)
    set_table_borders(meta, color=WHITE, size=0)
    values = (("文档版本", "V1.0"), ("编制日期", "2026年8月28日"), ("适用范围", "系统建设、交付与使用说明"))
    for row_idx, (label, value) in enumerate(values):
        set_cell_shading(meta.cell(row_idx, 0), LIGHT)
        set_cell_shading(meta.cell(row_idx, 1), LIGHT)
        set_cell_text(meta.cell(row_idx, 0), label, bold=True, color=TEAL, size=9.3,
                      align=WD_ALIGN_PARAGRAPH.RIGHT)
        set_cell_text(meta.cell(row_idx, 1), value, color=INK, size=9.3,
                      align=WD_ALIGN_PARAGRAPH.LEFT)
    doc.add_page_break()


def add_contents(doc: Document, title: str, items: Sequence[str]):
    doc.add_heading(title, level=1)
    add_callout(doc, "阅读提示", "本文档采用“结论先行、分层展开”的结构。各章节可独立查阅，也可按目录顺序完整阅读。")
    add_numbers(doc, items)
    doc.add_page_break()


def rounded_rectangle(draw: ImageDraw.ImageDraw, box, radius: int, fill: str, outline: str | None = None, width: int = 1):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def create_architecture_diagram(path: Path):
    img = Image.new("RGB", (1600, 1120), "#F7FAF9")
    draw = ImageDraw.Draw(img)
    title_font = ImageFont.truetype(FONT_FILE, 48)
    layer_font = ImageFont.truetype(FONT_FILE, 34)
    body_font = ImageFont.truetype(FONT_FILE, 24)
    small_font = ImageFont.truetype(FONT_FILE, 20)
    draw.text((800, 52), "视界学伴总体技术架构", font=title_font, fill="#163F3A", anchor="ma")
    layers = [
        ("交互呈现层", "首页 · 岗位分析 · 学习中心 · 实训工作台 · 模拟面试 · 成长档案", "#E5F0ED", "#2C6C62"),
        ("业务编排层", "生命周期状态机 · 意图路由 · 学习路径规划 · 自适应追问 · 角色协同", "#EDF3F6", "#456A7A"),
        ("领域模型层", "职业—岗位—任务—能力单元—能力点 · 前驱后继关系 · 评价证据", "#FBF3E3", "#C9973A"),
        ("数据资源层", "职业标准 · 教材原页 · 行业资料 · 岗位样本 · 实训案例 · 题库", "#F1EDEA", "#866B5A"),
        ("运行与交付层", "Vue 3 + TypeScript + Vite · 浏览器本地存储 · 单文件离线交付", "#E8EEF4", "#43607A"),
    ]
    y = 150
    for idx, (name, desc, fill, accent) in enumerate(layers):
        rounded_rectangle(draw, (120, y, 1480, y + 140), 24, fill, outline=accent, width=3)
        rounded_rectangle(draw, (150, y + 28, 445, y + 112), 18, accent)
        draw.text((297, y + 70), name, font=layer_font, fill="white", anchor="mm")
        draw.text((490, y + 49), desc, font=body_font, fill="#24313A")
        draw.text((490, y + 91), f"L{idx + 1}", font=small_font, fill=accent)
        if idx < len(layers) - 1:
            draw.line((800, y + 140, 800, y + 172), fill="#9DB1AB", width=5)
            draw.polygon([(790, y + 165), (810, y + 165), (800, y + 180)], fill="#9DB1AB")
        y += 188
    draw.text((800, 1080), "数据与能力关系贯穿各层，运行过程形成可追溯学习证据", font=body_font, fill="#66747D", anchor="mm")
    img.save(path, quality=95)


def create_lifecycle_diagram(path: Path):
    img = Image.new("RGB", (1800, 560), "#F7FAF9")
    draw = ImageDraw.Draw(img)
    title_font = ImageFont.truetype(FONT_FILE, 42)
    node_font = ImageFont.truetype(FONT_FILE, 28)
    small_font = ImageFont.truetype(FONT_FILE, 20)
    draw.text((900, 48), "岗位能力成长闭环", font=title_font, fill="#163F3A", anchor="ma")
    nodes = [
        ("多源资料", "职业标准 / 教材 / 岗位"),
        ("岗位解析", "任务 / 能力 / 关系"),
        ("学习转化", "理论 / 实训 / 路径"),
        ("能力验证", "检测 / 项目 / 面试"),
        ("成长反馈", "差距 / 档案 / 推荐"),
    ]
    xs = [180, 540, 900, 1260, 1620]
    fills = ["#E5F0ED", "#EDF3F6", "#FBF3E3", "#F1EDEA", "#E8EEF4"]
    accents = ["#2C6C62", "#456A7A", "#C9973A", "#866B5A", "#43607A"]
    for idx, ((name, desc), x) in enumerate(zip(nodes, xs)):
        rounded_rectangle(draw, (x - 145, 170, x + 145, 370), 28, fills[idx], outline=accents[idx], width=3)
        draw.ellipse((x - 34, 198, x + 34, 266), fill=accents[idx])
        draw.text((x, 232), str(idx + 1), font=node_font, fill="white", anchor="mm")
        draw.text((x, 300), name, font=node_font, fill="#24313A", anchor="mm")
        draw.text((x, 342), desc, font=small_font, fill="#66747D", anchor="mm")
        if idx < len(nodes) - 1:
            draw.line((x + 150, 270, xs[idx + 1] - 160, 270), fill="#9DB1AB", width=5)
            draw.polygon([(xs[idx + 1] - 174, 260), (xs[idx + 1] - 174, 280), (xs[idx + 1] - 155, 270)], fill="#9DB1AB")
    draw.arc((310, 360, 1490, 530), start=5, end=175, fill="#C9973A", width=4)
    draw.polygon([(318, 438), (340, 427), (336, 450)], fill="#C9973A")
    draw.text((900, 494), "验证结果回写能力状态，并驱动下一轮学习", font=small_font, fill="#866B5A", anchor="mm")
    img.save(path, quality=95)


def document_base(title: str, subject: str, short_title: str) -> Document:
    doc = Document()
    configure_styles(doc)
    configure_page(doc, short_title)
    doc.core_properties.title = title
    doc.core_properties.subject = subject
    doc.core_properties.author = "系统建设组"
    doc.core_properties.keywords = "视界学伴,计算机视觉,岗位能力,理实一体,智能助教"
    doc.core_properties.comments = "正式交付文档"
    return doc


def build_technology_document():
    architecture = WORK_DIR / "architecture.png"
    lifecycle = WORK_DIR / "lifecycle.png"
    create_architecture_diagram(architecture)
    create_lifecycle_diagram(lifecycle)

    doc = document_base(
        "视界学伴关键技术说明",
        "计算机视觉岗位能力成长系统的架构、核心机制与交付边界",
        "关键技术说明",
    )
    add_cover(doc, "关键技术说明", "计算机视觉岗位能力成长系统", "TECHNICAL DESCRIPTION")
    add_contents(doc, "文档目录", [
        "编制目的与技术定位", "总体技术架构", "多源岗位资料解析技术", "岗位能力图谱建模技术",
        "岗位任务到理实一体课程的转化", "多模态学习与智能助教", "多角色智能体协同实训",
        "自适应面试与薄弱点诊断", "学习证据与成长状态计算", "数据管理、离线交付与安全边界",
        "工程质量与扩展接口", "当前版本技术指标",
    ])

    doc.add_heading("一、编制目的与技术定位", level=1)
    add_paragraph(doc, "本文档说明视界学伴的总体技术架构、关键业务算法、数据组织方式、离线运行机制及工程边界，供系统建设、技术评审、部署交付和后续扩展使用。")
    add_callout(doc, "技术定位", "系统以岗位能力成长为主线，将多源岗位资料、能力图谱、理实一体学习、项目实训、模拟面试和成长档案组织为统一闭环。当前交付版本采用前端本地化运行方式，重点验证业务流程、领域模型与人机协作机制。")
    add_bullets(doc, [
        "岗位驱动：先识别岗位及典型工作任务，再生成课程与评价活动。",
        "证据驱动：阅读、检测、实训和面试分别形成不同效力等级的学习证据。",
        "可追溯：职业、岗位、任务、能力单元、能力点和课程资源之间保留明确关系。",
        "可离线：核心页面、教材图片、业务数据与交互逻辑可封装为单文件运行。",
    ])

    doc.add_heading("二、总体技术架构", level=1)
    add_picture(doc, architecture, "图 1  视界学伴总体技术架构", "视界学伴五层总体技术架构图", width=6.25)
    add_paragraph(doc, "系统采用分层前端架构。交互呈现层负责六类业务入口；业务编排层负责生命周期、意图、任务与评价流程；领域模型层沉淀岗位能力关系和学习证据；数据资源层承载标准、教材、岗位与案例；运行交付层提供类型校验、构建与离线封装能力。")
    add_table(doc, ["技术域", "主要实现", "承担职责", "技术价值"], [
        ["前端框架", "Vue 3 组合式 API", "组件化页面与响应式状态", "支撑复杂工作台的独立演进"],
        ["开发语言", "TypeScript", "领域类型、接口和状态约束", "减少跨模块数据漂移"],
        ["构建工具", "Vite + vue-tsc", "开发服务、类型检查与生产构建", "兼顾开发效率与交付稳定性"],
        ["图形呈现", "关系画布与 ECharts 能力", "岗位图谱、路径与数据可视化", "强化多层关系理解"],
        ["状态持久化", "浏览器 localStorage", "保存学习记录、分析状态和用户记忆", "无需后端即可恢复会话"],
        ["测试体系", "Vitest + jsdom", "领域算法、组件交互和离线入口校验", "形成可重复验证的工程基线"],
    ], [1500, 2100, 3060, 2700])

    doc.add_heading("三、多源岗位资料解析技术", level=1)
    add_paragraph(doc, "岗位分析工作台将教材、职业标准、行业分析报告和岗位说明材料作为可配置输入，并与内置产业、专业和招聘数据交叉归纳。解析过程采用八步可视化管线，使岗位结构的形成过程对使用者可见。")
    add_numbers(doc, [
        "读取教材与职业标准，校验资料类型、版本和来源标记。",
        "获取招聘岗位样本，载入岗位职责、技能和任职要求。",
        "识别岗位工作领域，归纳需求、数据、模型、部署和交付领域。",
        "提取典型工作任务，形成任务名称、步骤、产出和质量要求。",
        "拆解岗位能力单元，建立任务与能力单元映射。",
        "提取知识、技能与职业素养能力点。",
        "建立任务和能力点的前驱后继关系。",
        "组装岗位能力图谱，并附带来源和版本说明。",
    ])
    add_callout(doc, "可解释设计", "系统不直接展示一个无法复核的综合结论，而是同步呈现数据来源、处理阶段、结构数量和生成进度，使岗位分析过程具备可观察性。", fill=PALE_GOLD, accent=GOLD)

    doc.add_heading("四、岗位能力图谱建模技术", level=1)
    add_paragraph(doc, "岗位能力图谱采用“正式职业—职业方向—岗位—典型工作任务—能力单元—能力点”的层次模型。能力点进一步区分知识、技能和职业素养，并通过前驱后继关系连接学习顺序。")
    add_table(doc, ["层级", "对象说明", "典型属性", "主要关系"], [
        ["正式职业", "国家职业分类中的规范职业", "职业名称、目录编码、版本", "包含职业方向"],
        ["职业方向", "职业标准中的专业方向", "方向名称、标准锚点", "对应目标岗位"],
        ["典型工作任务", "岗位中可交付、可评价的任务", "步骤、产出、质量要求", "拆解能力单元"],
        ["能力单元", "完成任务所需的能力集合", "所属等级、任务归属", "聚合能力点"],
        ["能力点", "可学习、可观察、可验证的能力", "类型、描述、课程目标", "依赖前置、关联证据"],
    ], [1500, 2760, 2520, 2580])
    add_paragraph(doc, "图谱视图支持等级筛选、节点搜索、缩放、全量关系查看和节点详情检查。课程资源仅作为能力点的关联资源，不混入图谱层级，避免岗位标准结构与教学资源结构相互替代。")

    doc.add_heading("五、岗位任务到理实一体课程的转化", level=1)
    add_paragraph(doc, "系统将企业岗位任务转化为学习型工作任务，而不是直接把岗位职责复制为课程章节。转化过程依次生成学习目标、学习阶段、课程资源、实训情境、成果要求和评价证据。")
    add_bullets(doc, [
        "输入：典型工作任务、能力点、任务依赖、教材资源和实训案例。",
        "处理：知识、技能、素养目标拆解；理论与实践阶段编排；资源挂载。",
        "输出：学习型工作任务、前置与后继任务、课程页、实训项目和评价要求。",
        "门控：岗位分析完成后方可生成学习任务；未生成课程时不展示派生学习结果。",
    ])
    add_picture(doc, lifecycle, "图 2  岗位能力成长闭环", "岗位资料到成长反馈的五阶段闭环流程图", width=6.25)

    doc.add_heading("六、多模态学习与智能助教", level=1)
    add_paragraph(doc, "学习中心以真实教材原页为主体，配套章节目录、阅读进度、缩放控制、伴读提要、章节检测和页边笔记。教材图片、文字锚点、岗位能力和学习记录共同构成助教上下文。")
    add_table(doc, ["交互形态", "输入上下文", "系统输出", "证据效力"], [
        ["教材原页伴读", "当前页、章节、阅读锚点", "概念解释、岗位案例、启发式问题", "记录为学习中"],
        ["章节检测", "题目、答案、作答用时", "结果反馈与针对性学习陪伴", "通过后进入待验证"],
        ["岗位关联", "当前教材页与能力映射", "流式解析过程和能力卡片", "提供关系依据"],
        ["学习路径", "目标能力、前置关系、已有证据", "待学习前置、目标及后继能力", "驱动后续学习"],
    ], [1560, 2520, 3300, 1980])
    add_paragraph(doc, "当前助教回复采用本地规则、结构化映射和逐字流式呈现机制，能够展示上下文感知与交互闭环；未调用外部在线模型服务。")

    doc.add_heading("七、多角色智能体协同实训", level=1)
    add_paragraph(doc, "实训项目将工作过程划分为需求分析、数据处理、算法开发、测试验证和部署交付五类角色。学生选择一个主责角色，其余角色由协同智能体承担，并保留查看、退回、重做、批准和合并等协作动作。")
    add_table(doc, ["角色", "核心任务", "代表性交付物"], [
        ["需求分析", "澄清业务目标、验收规则与异常边界", "任务验收单.md"],
        ["数据处理", "样本检查、标注规范与数据集划分", "dataset-report.json"],
        ["算法开发", "基线模型、训练参数与推理流程", "train-config.yaml"],
        ["测试验证", "精度、召回率、时延与异常样本测试", "evaluation-report.md"],
        ["部署交付", "环境清单、部署脚本与回滚说明", "deployment-guide.md"],
    ], [1800, 4860, 2700])
    add_paragraph(doc, "协同状态采用 working、ready、approved、reworking 和 merged 等明确状态表达，能够呈现学生主责、智能体产出和人工审核之间的责任边界。")

    doc.add_heading("八、自适应面试与薄弱点诊断", level=1)
    add_paragraph(doc, "模拟面试设置算法、工程实践和系统设计三类面试官，可选择专项模式或联合轮换。每道题绑定能力单元、知识或技能评价要点，作答后按命中项与缺失项形成诊断。")
    add_bullets(doc, [
        "未覆盖关键要点时，下一题继续围绕当前能力域进行补强追问。",
        "覆盖核心要点时，下一题提升追问深度，增加量化指标、方案取舍、风险和验证证据要求。",
        "面试教练提供答题结构和示范建议，但不参与评分，避免辅导提示与评价结果混淆。",
        "诊断结果回写能力单元状态，并关联对应教材页和学习任务。",
    ])

    doc.add_heading("九、学习证据与成长状态计算", level=1)
    add_paragraph(doc, "系统将教材阅读、章节检测、实训项目和模拟面试统一抽象为学习证据。证据记录来源、对象、关联能力、结果和时间，并按效力计算能力状态。")
    add_table(doc, ["能力状态", "触发条件", "业务含义"], [
        ["未学习", "无关联证据", "尚未进入该能力的学习活动"],
        ["学习中", "仅有教材或课件访问证据", "已接触内容，但尚未接受评价"],
        ["待验证", "章节检测通过", "理论理解初步达成，仍需实践或面试验证"],
        ["已验证", "实训或面试通过", "形成较高效力的能力证据"],
        ["需要巩固", "最近一次评价未通过", "需要返回对应资源进行补强"],
    ], [1800, 3300, 4260])
    add_paragraph(doc, "学习路径规划采用能力依赖图的前驱遍历：先识别已满足前置，再输出未满足前置、目标能力和后继能力，并寻找关联学习任务。该机制可防止路径只按页面顺序排列。")

    doc.add_heading("十、数据管理、离线交付与安全边界", level=1)
    add_bullets(doc, [
        "会话持久化：采用版本化快照保存分析状态、学习证据、助教消息、成长记忆和面试记录。",
        "兼容迁移：当前学习会话可读取旧版本数据，并对不兼容字段执行回退或筛选。",
        "异常容错：浏览器存储不可用或数据解析失败时，保持内存会话并回退到默认状态。",
        "离线封装：构建脚本将 CSS、JavaScript 和图片资源内联到单一 HTML 文件。",
        "文件边界：当前本地文件选择主要读取文件名并建立材料清单，不向外部服务上传内容。",
    ])
    add_callout(doc, "当前技术边界", "系统未连接在线大模型接口、向量数据库、真实招聘接口或 GPU 训练环境；岗位解析、助教响应、协同智能体与面试自适应主要由本地结构化数据、状态机和规则逻辑驱动。", fill=PALE_GOLD, accent=GOLD)

    doc.add_heading("十一、工程质量与扩展接口", level=1)
    add_paragraph(doc, "系统通过领域类型、纯函数算法、组件测试和离线入口测试降低复杂交互的回归风险。主要扩展点如下：")
    add_bullets(doc, [
        "将岗位资料解析器替换为后端文档解析、检索增强和模型推理服务。",
        "将本地岗位样本替换为经授权的招聘数据服务，并保留来源、时间和数据批次。",
        "将本地存储替换为用户账户、学习记录存储和机构级数据治理服务。",
        "将规则型智能体替换为具备工具调用、权限控制和过程审计的智能体运行平台。",
        "将能力证据接口对接学习平台、实训平台、证书系统和就业服务系统。",
    ])

    doc.add_heading("十二、当前版本技术指标", level=1)
    add_table(doc, ["指标类别", "当前实现"], [
        ["业务主链", "岗位分析—学习与实训—模拟面试—成长档案"],
        ["岗位结构样例", "5 项典型工作任务、20 个能力单元、70 个能力点"],
        ["教材资源样例", "36 个教材原页学习节点"],
        ["实训资源样例", "4 个校内实训案例、5 类协同角色"],
        ["运行方式", "浏览器本地运行，支持开发入口与单文件离线入口"],
        ["数据外发", "当前版本不主动发起网络请求，不上传本地资料"],
    ], [2400, 6960])
    add_paragraph(doc, "以上指标为当前交付版本的系统配置，用于说明现有能力范围；后续可随专业、岗位、课程和机构数据扩展。", after=0)

    doc.save(TECH_DOC)


def build_function_document():
    doc = document_base(
        "视界学伴功能说明书",
        "计算机视觉岗位能力成长系统的功能、操作流程与业务规则",
        "功能说明书",
    )
    add_cover(doc, "功能说明书", "计算机视觉岗位能力成长系统", "FUNCTIONAL DESCRIPTION")
    add_contents(doc, "文档目录", [
        "系统概述", "适用对象与使用流程", "系统首页", "岗位分析资料配置", "岗位智能分析过程",
        "岗位能力图谱", "学习型工作任务与学习中心", "多角色协同实训", "实训项目工作台",
        "多智能体模拟面试", "个人成长档案", "业务规则与功能边界", "功能验收要点",
    ])

    doc.add_heading("一、系统概述", level=1)
    add_paragraph(doc, "视界学伴面向计算机视觉相关专业学习者，以目标岗位为起点，组织岗位解析、能力建模、课程学习、项目实训、模拟面试和成长反馈。系统强调“先理解岗位，再组织学习，最后用证据验证能力”。")
    add_callout(doc, "核心价值", "将职业标准、教材、岗位资料和实训资源转化为学生可理解、可执行、可验证的成长路径，并把各环节形成的证据统一沉淀到个人成长档案。")
    add_table(doc, ["功能域", "主要功能", "主要输出"], [
        ["岗位分析", "多源资料配置、八步解析、任务与能力拆解", "岗位能力图谱"],
        ["学习中心", "教材原页、智能伴读、检测、路径规划", "学习记录与待验证能力"],
        ["项目实训", "案例选择、角色协同、分步交付、产物审核", "实践证据与项目产物"],
        ["模拟面试", "多面试官、自适应追问、命中与缺失诊断", "面试记录与补强建议"],
        ["成长档案", "个人记忆、证据时间线、能力状态、历史记录", "持续更新的成长画像"],
    ], [1800, 4380, 3180])

    doc.add_heading("二、适用对象与使用流程", level=1)
    add_paragraph(doc, "主要使用对象包括专业学生、课程教师、实训指导教师和专业建设人员。学生负责完成学习、实训与面试；教师可使用岗位图谱、资源关联和证据记录组织教学与指导。")
    add_numbers(doc, [
        "进入首页，查看岗位能力成长闭环和当前进度。",
        "配置教材、职业标准、行业报告和岗位资料，运行岗位分析。",
        "查看典型工作任务、能力单元、能力点及其关系。",
        "将岗位任务转化为学习型工作任务，进入教材和实训学习。",
        "在实训项目中选择主责角色，并与其他角色智能体协同完成交付。",
        "参加多角色模拟面试，获得能力诊断和学习补强建议。",
        "在个人成长档案中复盘证据、能力状态和后续学习方向。",
    ])

    doc.add_heading("三、系统首页", level=1)
    add_paragraph(doc, "首页用于说明系统定位、展示完整业务闭环，并根据当前生命周期状态引导用户进入下一阶段。未完成岗位分析时，学习、面试等派生功能保持门控状态。")
    add_picture(doc, SHOT_DIR / "01-首页.png", "图 1  系统首页", "视界学伴系统首页，展示岗位能力成长闭环与开始岗位分析入口")
    add_bullets(doc, [
        "“开始岗位分析”进入资料配置工作台。",
        "“当前旅程”显示用户所处阶段和下一步建议。",
        "七个业务环节说明从资料接入到证书与就业推荐的完整路径。",
    ])

    doc.add_heading("四、岗位分析资料配置", level=1)
    add_paragraph(doc, "岗位分析支持教材、职业标准、行业分析报告和岗位说明书四类材料。系统预置示例资料，也允许用户选择本地文件建立材料清单。")
    add_picture(doc, SHOT_DIR / "02-岗位分析资料配置.png", "图 2  岗位分析资料配置", "岗位分析资料配置页，展示四类材料与已挂载资料")
    add_numbers(doc, [
        "检查已挂载资料的名称、类别和来源标记。",
        "如需补充资料，点击相应材料类型并选择本地文件。",
        "确认材料数量后，点击“开始智能分析”。",
    ])
    add_callout(doc, "文件处理说明", "当前版本建立本地资料清单并读取文件名，不向外部服务上传文件内容。预置资料用于保证系统在无外部依赖时能够完整运行。", fill=PALE_GOLD, accent=GOLD)

    doc.add_heading("五、岗位智能分析过程", level=1)
    add_paragraph(doc, "运行分析后，系统展示八步解析过程、数据接入状态、岗位样本数量、典型任务数量、能力节点数量和关系数量。用户能够看到岗位结构从资料到图谱的形成过程。")
    add_picture(doc, SHOT_DIR / "03-岗位智能分析过程.png", "图 3  岗位智能分析过程", "岗位智能分析运行页，展示数据来源、进度和分步解析状态")
    add_bullets(doc, [
        "顶部显示当前解析目标和岗位数据更新时间。",
        "中部显示产业、专业与招聘数据的接入状态。",
        "底部按顺序展示资料读取、岗位识别、任务提取、能力拆解和图谱生成。",
        "分析完成后自动进入岗位能力图谱。",
    ])

    doc.add_heading("六、岗位能力图谱", level=1)
    add_paragraph(doc, "岗位能力图谱是系统的核心关系视图。页面同时展示职业与岗位上下文、典型工作任务数量、能力单元数量、能力点数量和已转学习任务数量。")
    add_picture(doc, SHOT_DIR / "04-岗位能力图谱.png", "图 4  岗位能力图谱", "岗位能力图谱页面，展示职业、岗位、任务、能力单元和能力点关系")
    add_table(doc, ["操作", "功能说明"], [
        ["等级切换", "按初级、中级、高级查看对应岗位任务和能力要求"],
        ["节点搜索", "按能力单元或能力点名称快速定位"],
        ["图谱缩放", "调整复杂关系的可见范围"],
        ["查看全量", "恢复并展示完整任务、能力单元和能力点关系"],
        ["节点详情", "查看对象说明、标准依据、前驱后继和课程关联"],
        ["批量生成课程", "将当前等级典型工作任务转化为理实一体学习任务"],
    ], [2100, 7260])
    add_paragraph(doc, "图谱中的职业编码和职业方向用于说明岗位的规范来源；岗位名称用于教学与就业场景表达。课程资源是能力点的关联资源，不属于职业标准原有层级。")

    doc.add_heading("七、学习型工作任务与学习中心", level=1)
    add_paragraph(doc, "用户可在岗位图谱中批量生成学习型工作任务。生成完成后，学习中心按岗位任务组织教材原页、课程目标、前后置任务、实训项目和评价证据。")
    add_picture(doc, SHOT_DIR / "05-理实一体学习中心.png", "图 5  理实一体学习中心", "理实一体学习中心，包含教材目录、真实教材原页和智能助教")
    add_bullets(doc, [
        "左侧目录按任务或教材章节组织真实教材原页，并显示阅读进度。",
        "中部阅读区提供教材页码、来源、缩放、伴读提要、章节检测和笔记。",
        "右侧智能助教感知当前教材页、岗位能力和学习记录，可进行逐页伴读、岗位举例、启发式提问、岗位关联和学习路径规划。",
        "章节检测记录作答结果与用时，并根据表现给出针对性学习陪伴。",
    ])

    doc.add_heading("八、多角色协同实训", level=1)
    add_paragraph(doc, "进入实训后，学生首先选择自己的主责角色。其余四个角色由智能体团队执行，学生可以查看、审核、退回或合并智能体产物。")
    add_picture(doc, SHOT_DIR / "06-多角色协同实训.png", "图 6  多角色协同实训角色选择", "工业表面缺陷检测实训的五类角色选择页面")
    add_table(doc, ["角色", "职责", "主责产物"], [
        ["需求分析", "业务目标、验收规则、异常边界", "任务验收单"],
        ["数据处理", "样本检查、标注规范、数据划分", "数据报告"],
        ["算法开发", "模型配置、训练参数、推理流程", "训练配置"],
        ["测试验证", "精度、召回率、时延和异常样本", "评价报告"],
        ["部署交付", "环境、部署、操作与回滚", "部署指南"],
    ], [1800, 4560, 3000])
    add_paragraph(doc, "角色可在项目过程中切换。切换后，新的主责角色由学生执行，原学生角色转为智能体协同角色，保证职责分工清晰。")

    doc.add_heading("九、实训项目工作台", level=1)
    add_paragraph(doc, "实训工作台以真实项目交付顺序组织任务。当前示例“工业表面缺陷检测”设置需求分析、算法设计、测试验证和部署优化四个步骤，前一步提交后解锁下一步。")
    add_picture(doc, SHOT_DIR / "07-实训项目工作台.png", "图 7  实训项目工作台", "工业表面缺陷检测实训工作台，展示步骤、资源和交付表单")
    add_bullets(doc, [
        "项目顶部显示业务场景、关联岗位任务、质量指标和当前主责角色。",
        "左侧显示步骤进度和已加载的代码、数据、评测与优秀项目资源。",
        "中部按步骤提供结构化交付表单，避免只提交笼统结论。",
        "支持选择本地产物或载入实训样例；样例必须由学生核验后提交。",
        "智能体产物需要经过人工审核，可执行批准、退回重做或合并。",
    ])

    doc.add_heading("十、多智能体模拟面试", level=1)
    add_paragraph(doc, "模拟面试通过算法、工程实践和系统设计三类面试官验证目标等级能力。联合面试模式下，面试官按题轮换；专项模式下，围绕某一能力域持续追问。")
    add_picture(doc, SHOT_DIR / "08-多智能体模拟面试.png", "图 8  多智能体模拟面试", "岗位模拟面试页面，展示面试官阵容、题目和面试教练")
    add_numbers(doc, [
        "选择面试官模式和目标岗位等级。",
        "阅读题目、题型、考察重点和作答要求。",
        "按业务目标、技术方案、量化指标、风险和证据组织回答。",
        "提交回答后查看命中要点、缺失要点和下一题自适应方向。",
        "面试结束后查看能力诊断、对应教材和后续学习建议。",
    ])
    add_callout(doc, "评分边界", "“智能助教·面试教练”只提供答题结构、示范和改进建议，不参与评分。评分依据来自面试题绑定的评价要点及用户实际回答。", fill=PALE_GOLD, accent=GOLD)

    doc.add_heading("十一、个人成长档案", level=1)
    add_paragraph(doc, "个人成长档案统一沉淀教材进度、有效证据、能力状态、个人记忆、面试记录和证据时间线。用户可从档案继续学习路径或再次参加面试。")
    add_picture(doc, SHOT_DIR / "09-个人成长档案.png", "图 9  个人成长档案", "个人成长档案页面，展示学习指标、智能体记忆和学习证据时间线")
    add_table(doc, ["模块", "功能说明"], [
        ["成长指标", "显示教材进度、有效证据、已验证能力和待验证能力"],
        ["个人记忆", "记录目标岗位、近期重点和学习偏好，支持编辑与删除"],
        ["历史面试", "保存整场面试记录、诊断和岗位建议"],
        ["证据时间线", "按时间展示阅读、检测、实训和面试证据"],
        ["能力状态", "区分未学习、学习中、待验证、已验证和需要巩固"],
    ], [2100, 7260])

    doc.add_heading("十二、业务规则与功能边界", level=1)
    add_table(doc, ["规则", "说明"], [
        ["生命周期门控", "未完成岗位分析时不生成岗位、课程、面试等派生结果"],
        ["课程生成门控", "课程由典型工作任务转化生成，不提供脱离岗位任务的默认课程结果"],
        ["阅读不等于掌握", "教材访问只记录学习中，不能直接把能力标记为已验证"],
        ["高效力证据", "实训或面试通过后，相关能力才进入已验证状态"],
        ["证据优先", "系统保留岗位、任务、能力和来源，不展示无法可靠证明的岗位匹配比例"],
        ["本地运行", "当前版本不调用在线大模型、真实招聘接口或 GPU 训练环境"],
    ], [2400, 6960])
    add_paragraph(doc, "系统中的岗位名称、课程转化结果和实训任务用于教学场景表达。正式职业、职业方向和职业编码以对应版本的职业标准与职业分类资料为依据。")

    doc.add_heading("十三、功能验收要点", level=1)
    add_bullets(doc, [
        "首次进入系统时，仅显示首页和岗位分析入口，不提前泄漏派生结果。",
        "岗位分析能够展示资料配置、运行进度和分析完成后的能力图谱。",
        "图谱能够切换等级、搜索节点、查看任务—能力单元—能力点关系。",
        "学习任务必须由岗位任务转化生成，生成后可进入教材和实训资源。",
        "教材原页、智能助教、章节检测和学习证据之间能够形成闭环。",
        "实训能够选择主责角色，展示其他角色智能体并支持审核动作。",
        "模拟面试能够切换面试官、记录回答、生成诊断和补强建议。",
        "成长档案能够恢复学习进度、证据、记忆和历史面试记录。",
        "开发入口与单文件离线入口均可正常打开并完成关键流程。",
    ])
    add_callout(doc, "交付结论", "视界学伴已形成岗位分析、学习与实训、模拟面试和成长档案相互连接的完整功能链。当前版本适用于本地展示、教学研讨、功能验证和后续系统化建设。")

    doc.save(FUNC_DOC)


if __name__ == "__main__":
    build_technology_document()
    build_function_document()
    print(TECH_DOC)
    print(FUNC_DOC)
