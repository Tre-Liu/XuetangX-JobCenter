from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
SOURCE = ROOT / "docs/superpowers/specs/2026-06-02-product-implementation-phases-design.md"
OUTPUT = ROOT / "outputs/prd/产品实现规划文档.docx"

TITLE = "产品实现规划文档"
SUBTITLE = "专业建设平台分期建设与功能落地规划"
META = "适用对象：产品 / 研发    文档版本：V1.0    日期：2026-06-02"

BODY_FONT = "Songti SC"
HEADING_FONT = "PingFang SC"
LATIN_FONT = "Calibri"

INK = RGBColor(36, 53, 84)
BLUE = RGBColor(46, 116, 181)
DARK_BLUE = RGBColor(31, 77, 120)
MUTED = RGBColor(96, 108, 125)

HEADING_STYLES = {
    1: ("Heading 1", 16, BLUE, 16, 8),
    2: ("Heading 2", 13, BLUE, 12, 6),
    3: ("Heading 3", 12, DARK_BLUE, 8, 4),
    4: ("Heading 4", 11.5, INK, 8, 3),
    5: ("Heading 5", 11, INK, 6, 3),
}


def set_east_asia_font(target, east_asia: str, ascii_font: str | None = None) -> None:
    if ascii_font is None:
        ascii_font = east_asia
    target.font.name = ascii_font
    rpr = target._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    rfonts.set(qn("w:ascii"), ascii_font)
    rfonts.set(qn("w:hAnsi"), ascii_font)
    rfonts.set(qn("w:eastAsia"), east_asia)


def add_page_number(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def add_bottom_rule(paragraph, color: str = "D9E1F2", size: str = "8") -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), size)
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    p_bdr.append(bottom)


def set_document_geometry(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)


def configure_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    set_east_asia_font(normal, BODY_FONT, LATIN_FONT)
    normal.font.size = Pt(11)
    normal.font.color.rgb = INK
    pf = normal.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after = Pt(6)
    pf.line_spacing = 1.1

    for _, (style_name, size, color, before, after) in HEADING_STYLES.items():
        style = document.styles[style_name]
        set_east_asia_font(style, HEADING_FONT, LATIN_FONT)
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for style_name in ["List Number", "List Bullet"]:
        style = document.styles[style_name]
        set_east_asia_font(style, BODY_FONT, LATIN_FONT)
        style.font.size = Pt(11)
        style.font.color.rgb = INK
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.15


def set_running_header_footer(section) -> None:
    header = section.header
    header_p = header.paragraphs[0]
    header_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    header_p.paragraph_format.space_after = Pt(0)
    header_run = header_p.add_run(TITLE)
    set_east_asia_font(header_run, BODY_FONT, LATIN_FONT)
    header_run.font.size = Pt(9)
    header_run.font.color.rgb = MUTED

    footer = section.footer
    footer_p = footer.paragraphs[0]
    footer_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    footer_p.paragraph_format.space_after = Pt(0)
    label = footer_p.add_run("第 ")
    set_east_asia_font(label, BODY_FONT, LATIN_FONT)
    label.font.size = Pt(9)
    label.font.color.rgb = MUTED
    add_page_number(footer_p)
    tail = footer_p.add_run(" 页")
    set_east_asia_font(tail, BODY_FONT, LATIN_FONT)
    tail.font.size = Pt(9)
    tail.font.color.rgb = MUTED


def add_cover(document: Document) -> None:
    for _ in range(3):
        document.add_paragraph()

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    run = p.add_run(TITLE)
    set_east_asia_font(run, HEADING_FONT, LATIN_FONT)
    run.font.size = Pt(24)
    run.font.bold = True
    run.font.color.rgb = INK

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(SUBTITLE)
    set_east_asia_font(run, BODY_FONT, LATIN_FONT)
    run.font.size = Pt(12)
    run.font.color.rgb = MUTED

    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(18)
    run = p.add_run(META)
    set_east_asia_font(run, BODY_FONT, LATIN_FONT)
    run.font.size = Pt(10.5)
    run.font.color.rgb = MUTED

    rule = document.add_paragraph()
    rule.paragraph_format.space_after = Pt(18)
    add_bottom_rule(rule)

    intro = document.add_paragraph()
    intro.paragraph_format.space_after = Pt(14)
    intro.alignment = WD_ALIGN_PARAGRAPH.LEFT
    intro_run = intro.add_run(
        "本文档用于指导专业建设平台的产品与研发分期落地，围绕岗位、产业、课程和人才培养方案四类核心对象，明确建设目标、能力边界与阶段承接关系。"
    )
    set_east_asia_font(intro_run, BODY_FONT, LATIN_FONT)
    intro_run.font.size = Pt(11)
    intro_run.font.color.rgb = INK

    document.add_section(WD_SECTION_START.NEW_PAGE)


def clean_text(text: str) -> str:
    return re.sub(r"\s+", " ", text.replace("  ", " ")).strip()


def append_paragraph(document: Document, text: str) -> None:
    if not text:
        return
    p = document.add_paragraph(style="Normal")
    p.paragraph_format.first_line_indent = Inches(0.28)
    run = p.add_run(text)
    set_east_asia_font(run, BODY_FONT, LATIN_FONT)
    run.font.size = Pt(11)
    run.font.color.rgb = INK


def append_heading(document: Document, level: int, text: str) -> None:
    level = max(1, min(5, level))
    style_name = HEADING_STYLES[level][0]
    p = document.add_paragraph(style=style_name)
    run = p.add_run(text.strip())
    set_east_asia_font(run, HEADING_FONT, LATIN_FONT)
    run.font.bold = True


def append_number_list(document: Document, text: str) -> None:
    p = document.add_paragraph(style="List Number")
    p.paragraph_format.first_line_indent = Inches(0)
    p.paragraph_format.left_indent = Inches(0.25)
    run = p.add_run(text)
    set_east_asia_font(run, BODY_FONT, LATIN_FONT)
    run.font.size = Pt(11)
    run.font.color.rgb = INK


def append_bullet_list(document: Document, text: str) -> None:
    p = document.add_paragraph(style="List Bullet")
    p.paragraph_format.first_line_indent = Inches(0)
    p.paragraph_format.left_indent = Inches(0.25)
    run = p.add_run(text)
    set_east_asia_font(run, BODY_FONT, LATIN_FONT)
    run.font.size = Pt(11)
    run.font.color.rgb = INK


def build_document() -> None:
    document = Document()
    set_document_geometry(document)
    configure_styles(document)
    document.core_properties.title = TITLE
    document.core_properties.subject = "专业建设平台分期规划"
    document.core_properties.author = "OpenAI Codex"
    document.core_properties.comments = "Converted from approved planning markdown."

    set_running_header_footer(document.sections[0])
    add_cover(document)
    set_running_header_footer(document.sections[-1])

    lines = SOURCE.read_text(encoding="utf-8").splitlines()
    i = 0
    paragraph_buffer: list[str] = []

    def flush_buffer() -> None:
        if paragraph_buffer:
            append_paragraph(document, clean_text(" ".join(paragraph_buffer)))
            paragraph_buffer.clear()

    while i < len(lines):
        raw = lines[i]
        stripped = raw.strip()

        if not stripped:
            flush_buffer()
            i += 1
            continue

        if stripped == "# 产品实现规划文档":
            i += 1
            continue

        heading_match = re.match(r"^(#{2,6})\s+(.*)$", stripped)
        if heading_match:
            flush_buffer()
            level = len(heading_match.group(1)) - 1
            append_heading(document, level, heading_match.group(2))
            i += 1
            continue

        number_match = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if number_match:
            flush_buffer()
            item_text = number_match.group(2).strip()
            i += 1
            while i < len(lines):
                nxt = lines[i]
                nxt_stripped = nxt.strip()
                if not nxt_stripped:
                    i += 1
                    break
                if re.match(r"^(#{2,6})\s+", nxt_stripped):
                    break
                if re.match(r"^\d+\.\s+", nxt_stripped):
                    break
                if re.match(r"^- ", nxt_stripped):
                    break
                item_text += " " + nxt_stripped
                i += 1
            append_number_list(document, clean_text(item_text))
            continue

        bullet_match = re.match(r"^- (.*)$", stripped)
        if bullet_match:
            flush_buffer()
            item_text = bullet_match.group(1).strip()
            i += 1
            while i < len(lines):
                nxt = lines[i]
                nxt_stripped = nxt.strip()
                if not nxt_stripped:
                    i += 1
                    break
                if re.match(r"^(#{2,6})\s+", nxt_stripped):
                    break
                if re.match(r"^\d+\.\s+", nxt_stripped):
                    break
                if re.match(r"^- ", nxt_stripped):
                    break
                item_text += " " + nxt_stripped
                i += 1
            append_bullet_list(document, clean_text(item_text))
            continue

        paragraph_buffer.append(stripped)
        i += 1

    flush_buffer()
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)


if __name__ == "__main__":
    build_document()
